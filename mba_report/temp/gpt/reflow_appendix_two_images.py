# -*- coding: utf-8 -*-
"""Rebuild the screenshot appendix with two images per page."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_DOC = BASE_DIR / "MBA_중간점검_과제물_이현용_부록포함.standard-docx"
OUTPUT_STANDARD = BASE_DIR / "MBA_중간점검_과제물_이현용_부록2장씩.standard-docx"
OUTPUT_DOCX = BASE_DIR / "MBA_중간점검_과제물_이현용_부록2장씩.docx"
PHOTO_DIR = BASE_DIR / "사진"

PHOTO_ORDER = [
    ("로그인", "로그인.png"),
    ("프로필 설정 - 나이, 성별, 무게, 키", "프로필 설정(나이_성별_무게_키).png"),
    ("프로필 설정 - 활동량", "프로필 설정(활동량).png"),
    ("프로필 설정 - 체중 목표", "프로필 설정(체중 목표).png"),
    ("프로필 설정 - 추천 칼로리 및 목표 설정", "프로필 설정(추천 칼로리_목표 설정).png"),
    ("홈", "홈.png"),
    ("기록", "기록.png"),
    ("통계", "통계.png"),
    ("설정", "설정.png"),
]


def set_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def clear_appendix(doc: Document) -> None:
    start = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "부록. 실제 화면 캡처":
            start = paragraph._element
            break
    if start is None:
        raise ValueError("Appendix heading not found.")

    body = doc._body._element
    children = list(body)
    start_index = children.index(start)
    for element in children[start_index:]:
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)


def add_centered_paragraph(cell, text: str, size: float, bold: bool = False):
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    return paragraph


def add_image_cell(cell, index: int, caption: str, filename: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.width = Inches(3.05)
    add_centered_paragraph(cell, f"부록 {index}. {caption}", size=9.5, bold=True)

    image_paragraph = cell.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = image_paragraph.add_run()
    image_run.add_picture(str(PHOTO_DIR / filename), width=Inches(2.72))

    filename_paragraph = cell.add_paragraph()
    filename_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    filename_run = filename_paragraph.add_run(filename)
    set_font(filename_run, size=7.5)


def rebuild_appendix() -> Path:
    missing = [filename for _, filename in PHOTO_ORDER if not (PHOTO_DIR / filename).exists()]
    if missing:
        raise FileNotFoundError("Missing appendix images: " + ", ".join(missing))

    doc = Document(SOURCE_DOC)
    clear_appendix(doc)

    doc.add_page_break()
    heading = doc.add_heading("", level=1)
    heading_run = heading.add_run("부록. 실제 화면 캡처")
    set_font(heading_run, size=16, bold=True)

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "아래 이미지는 Cal AI 앱의 실제 사용 흐름에 따라 정리했으며, 한 페이지에 두 장씩 배치하였다."
    )
    set_font(intro_run, size=10.5)

    for pair_start in range(0, len(PHOTO_ORDER), 2):
        if pair_start > 0:
            doc.add_page_break()

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)

        for col_index in range(2):
            cell = table.cell(0, col_index)
            cell.text = ""
            item_index = pair_start + col_index
            if item_index < len(PHOTO_ORDER):
                caption, filename = PHOTO_ORDER[item_index]
                add_image_cell(cell, item_index + 1, caption, filename)

    doc.save(OUTPUT_STANDARD)
    shutil.copyfile(OUTPUT_STANDARD, OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(rebuild_appendix())
