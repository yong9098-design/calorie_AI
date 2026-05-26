# -*- coding: utf-8 -*-
"""Append screenshot photos to the MBA report as an ordered appendix."""
from __future__ import annotations

from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


BASE_DIR = Path(__file__).resolve().parent
SOURCE_DOC = BASE_DIR / "MBA_중간점검_과제물_이현용.standard-docx"
OUTPUT_DOC = BASE_DIR / "MBA_중간점검_과제물_이현용_부록포함.standard-docx"
OUTPUT_DOCX = BASE_DIR / "MBA_중간점검_과제물_이현용_부록포함.docx"
PHOTO_DIR = BASE_DIR / "사진"

PHOTO_ORDER = [
    ("로그인", "로그인.png"),
    ("프로필 설정 - 나이, 성별, 무게, 키", "프로필 설정(나이_성별_무게_키).png"),
    ("프로필 설정 - 활동량", "프로필 설정(활동량).png"),
    ("프로필 설정 - 체중 목표", "프로필 설정(체중 목표).png"),
    ("프로필 설정 - 추천 칼로리 및 목표 설정", "프로필 설정(추천 칼로리_목표 설정).png"),
    ("홈", "홈.png"),
    ("기록", "기록.png"),
    ("통계", "통계.png"),
    ("설정", "설정.png"),
]


def set_korean_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def max_picture_width(path: Path) -> Inches:
    """Keep mobile screenshots readable without overflowing the Word page."""
    with Image.open(path) as img:
        width, height = img.size
    if height > width:
        return Inches(3.4)
    return Inches(6.1)


def add_centered_text(doc: Document, text: str, size: float, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_korean_font(run, size=size, bold=bold)


def add_photo_appendix() -> Path:
    missing = [filename for _, filename in PHOTO_ORDER if not (PHOTO_DIR / filename).exists()]
    if missing:
        raise FileNotFoundError("Missing appendix images: " + ", ".join(missing))

    doc = Document(SOURCE_DOC)
    doc.add_page_break()

    title = doc.add_heading("", level=1)
    title_run = title.add_run("부록. 실제 화면 캡처")
    set_korean_font(title_run, size=16, bold=True)

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "아래 이미지는 Cal AI 앱의 실제 사용 흐름에 따라 로그인, 프로필 설정, 주요 기능 화면 순서로 정리하였다."
    )
    set_korean_font(intro_run, size=10.5)

    for index, (caption, filename) in enumerate(PHOTO_ORDER, start=1):
        if index > 1:
            doc.add_page_break()

        add_centered_text(doc, f"부록 {index}. {caption}", size=11, bold=True)

        image_path = PHOTO_DIR / filename
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = picture.add_run()
        run.add_picture(str(image_path), width=max_picture_width(image_path))

        filename_line = doc.add_paragraph()
        filename_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        filename_run = filename_line.add_run(filename)
        set_korean_font(filename_run, size=9)

    doc.save(OUTPUT_DOC)
    copyfile(OUTPUT_DOC, OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(add_photo_appendix())
