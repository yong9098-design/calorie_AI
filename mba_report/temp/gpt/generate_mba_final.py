# -*- coding: utf-8 -*-
"""
MBA 중간점검 과제물 최종 생성 스크립트
- 1-1: 개인 다이어트 문단 추가 / MBA 이유 문단 삭제
- 1-2: 기존 테이블 삭제 / 앱 비용 관련 새 내용 추가
- 1-3: 아키텍처 다이어그램 개선 (겹침 해소, 고품질)
- 부록: 스크린샷 2장씩 배치
"""
from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


BASE_DIR = Path(__file__).resolve().parent
PHOTO_DIR = BASE_DIR / "사진"
ARCH_PATH = BASE_DIR / "temp" / "architecture_diagram_v2.png"
TECH_PATH = BASE_DIR / "temp" / "tech_stack_diagram_v2.png"
SCREEN_PATH = BASE_DIR / "temp" / "screen_list_diagram_v2.png"
OUTPUT_STANDARD = BASE_DIR / "MBA_중간점검_과제물_이현용_최종.standard-docx"
OUTPUT_DOCX = BASE_DIR / "MBA_중간점검_과제물_이현용_최종.docx"

PHOTO_ORDER = [
    ("로그인", "로그인.png"),
    ("프로필 설정 - 나이, 성별, 무게, 키", "프로필 설정(나이_성별_무게_키).png"),
    ("프로필 설정 - 활동량", "프로필 설정(활동량).png"),
    ("프로필 설정 - 체중 목표", "프로필 설정(체중 목표).png"),
    ("프로필 설정 - 추천 칼로리 및 목표 설정", "프로필 설정(추천 칼로리_목표 설정).png"),
    ("홈", "홈.png"),
    ("기록", "기록.png"),
    ("통계", "통계.png"),
    ("설정", "설정.png"),
]


# ─────────────────────────────────────────────────────────────────────────────
# Shared drawing helpers
# ─────────────────────────────────────────────────────────────────────────────

def setup_plot() -> None:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def rbox(ax, x, y, w, h, text, face, edge=None, tc="#ffffff", fs=12, lw=2.0):
    """Draw a rounded rectangle with centered text."""
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.04,rounding_size=0.09",
        linewidth=lw,
        edgecolor=edge or face,
        facecolor=face,
        zorder=3,
    )
    ax.add_patch(patch)
    ax.text(
        x + w / 2, y + h / 2, text,
        ha="center", va="center",
        color=tc, fontsize=fs, fontweight="bold",
        linespacing=1.35, zorder=4,
    )


def arr(ax, x0, y0, x1, y1, color="#2563eb", lw=2.2, rad=0.0):
    """Draw a clean arrow."""
    ax.annotate(
        "",
        xy=(x1, y1), xytext=(x0, y0),
        arrowprops={
            "arrowstyle": "->",
            "lw": lw,
            "color": color,
            "shrinkA": 4,
            "shrinkB": 4,
            "connectionstyle": f"arc3,rad={rad}",
        },
        zorder=5,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Diagram 1: Architecture (improved — no overlaps)
# ─────────────────────────────────────────────────────────────────────────────

def make_architecture() -> Path:
    """
    Layout (xlim=0..16, ylim=0..11):
      y=9.4  title
      y=8.0-8.7  doc boxes  (PRD.md, SPEC.md)
      y=5.8-7.1  main pipeline boxes
      y=3.3-4.1  feedback boxes (output/index.html, QA_REPORT.md)
      y=2.55     feedback label (red text, clearly below red arrows)
      y=2.0      dashed divider
      y=0.5-1.4  tech stack
    """
    setup_plot()
    fig, ax = plt.subplots(figsize=(16, 11))
    fig.patch.set_facecolor("#f0f4f8")
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis("off")

    # Title
    ax.text(8, 10.3, "Cal AI 3-Agent 파이프라인 아키텍처",
            ha="center", fontsize=22, fontweight="bold", color="#0f172a", zorder=4)

    # ── Main pipeline boxes ──────────────────────────────────────────────────
    Y_MAIN = 5.8
    H_BOX = 1.3
    Y_CENTER = Y_MAIN + H_BOX / 2  # 6.45

    # Box definitions: (x, width, label, face_color)
    main_boxes = [
        (0.20, 1.80, "사용자\n요청",  "#475569"),
        (2.60, 2.30, "Planner\nAgent", "#0284c7"),
        (5.80, 2.30, "Generator\nAgent", "#16a34a"),
        (9.00, 2.30, "Evaluator\nAgent", "#f59e0b"),
        (12.20, 2.30, "합격\n완료", "#22c55e"),
    ]
    for x, w, label, face in main_boxes:
        rbox(ax, x, Y_MAIN, w, H_BOX, label, face, fs=13)

    # Horizontal arrows between boxes
    gaps = [
        (2.00, 2.60),   # user → planner
        (4.90, 5.80),   # planner → generator
        (8.10, 9.00),   # generator → evaluator
        (11.30, 12.20), # evaluator → done
    ]
    for x0, x1 in gaps:
        arr(ax, x0, Y_CENTER, x1, Y_CENTER, "#2563eb", lw=2.4)

    # ── Document boxes above ─────────────────────────────────────────────────
    Y_DOC = 7.8
    H_DOC = 0.75

    # PRD.md above Planner
    rbox(ax, 2.70, Y_DOC, 2.10, H_DOC, "PRD.md",
         "#e0f2fe", edge="#0284c7", tc="#075985", fs=11, lw=1.8)
    # Arrow: PRD.md → Planner top
    arr(ax, 3.75, Y_DOC, 3.75, Y_MAIN + H_BOX, "#64748b", lw=2.0)

    # SPEC.md above Generator
    rbox(ax, 5.90, Y_DOC, 2.10, H_DOC, "SPEC.md",
         "#dcfce7", edge="#16a34a", tc="#166534", fs=11, lw=1.8)
    # Arrow: SPEC.md → Generator top
    arr(ax, 6.95, Y_DOC, 6.95, Y_MAIN + H_BOX, "#64748b", lw=2.0)

    # Planner → SPEC.md (shows Planner writes SPEC.md)
    arr(ax, 3.75, Y_DOC + H_DOC, 6.90, Y_DOC + H_DOC, "#64748b", lw=1.6)
    ax.text(5.30, Y_DOC + H_DOC + 0.12, "SPEC.md 생성",
            ha="center", va="bottom", fontsize=8.5, color="#475569", style="italic")

    # ── Feedback loop boxes ──────────────────────────────────────────────────
    Y_FB = 3.3
    H_FB = 0.85

    rbox(ax, 5.70, Y_FB, 2.40, H_FB, "output/\nindex.html",
         "#f8fafc", edge="#64748b", tc="#334155", fs=10.5, lw=1.8)

    rbox(ax, 8.80, Y_FB, 2.70, H_FB, "QA_REPORT.md\n점수 / 피드백",
         "#ffedd5", edge="#f59e0b", tc="#9a3412", fs=10, lw=1.8)

    # Evaluator → QA_REPORT (grey, downward)
    arr(ax, 10.15, Y_MAIN, 10.15, Y_FB + H_FB, "#64748b", lw=2.0)

    # QA_REPORT → output/index.html (red, leftward)
    arr(ax, 8.80, Y_FB + H_FB / 2, 8.10, Y_FB + H_FB / 2, "#ef4444", lw=2.2)

    # output/index.html → Generator (red, upward)
    arr(ax, 6.90, Y_FB + H_FB, 6.90, Y_MAIN, "#ef4444", lw=2.2)

    # Feedback label — clearly positioned below the red arrows with a light box
    label_y = 2.52
    label_x = 7.5
    label_text = "  불합격 또는 조건부 합격 시 피드백 반영  "
    bg = FancyBboxPatch((4.85, 2.32), 5.3, 0.42,
                        boxstyle="round,pad=0.04", linewidth=1.2,
                        edgecolor="#fca5a5", facecolor="#fff1f1", zorder=3)
    ax.add_patch(bg)
    ax.text(label_x, label_y, label_text,
            ha="center", va="center",
            fontsize=10.5, color="#dc2626", fontweight="bold", zorder=5)

    # ── Dashed divider ───────────────────────────────────────────────────────
    ax.plot([0.5, 15.5], [1.95, 1.95],
            linestyle=(0, (6, 4)), color="#94a3b8", linewidth=1.8, zorder=2)

    # ── Tech stack ───────────────────────────────────────────────────────────
    tech = [
        ("HTML/CSS/JS\n모바일 SPA", "#0ea5e9"),
        ("Gemini 2.5\n이미지 분석", "#7c3aed"),
        ("Supabase\nAuth + DB", "#059669"),
        ("Vercel\n배포 + API", "#111827"),
    ]
    tw, tgap = 2.8, 0.55
    tx_start = (16 - (len(tech) * tw + (len(tech) - 1) * tgap)) / 2
    for i, (label, color) in enumerate(tech):
        tx = tx_start + i * (tw + tgap)
        rbox(ax, tx, 0.45, tw, 1.1, label, color, fs=11)

    ARCH_PATH.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(ARCH_PATH, dpi=220, bbox_inches="tight",
                facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)
    return ARCH_PATH


# ─────────────────────────────────────────────────────────────────────────────
# Diagram 2: Tech Stack
# ─────────────────────────────────────────────────────────────────────────────

def make_tech_stack() -> Path:
    setup_plot()
    fig, ax = plt.subplots(figsize=(14, 8.5))
    fig.patch.set_facecolor("#ffffff")
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8.5)
    ax.axis("off")
    ax.text(7, 8.1, "Cal AI 기술 스택", ha="center",
            fontsize=22, fontweight="bold", color="#0f172a")

    layers = [
        ("사용자 화면", 5.5, "#dbeafe", "#1d4ed8",
         [("HTML5", "구조"), ("CSS3", "모바일 UI"), ("Vanilla JS", "SPA 상태"), ("SVG/Canvas", "차트")]),
        ("AI Agent 개발", 3.2, "#dcfce7", "#15803d",
         [("Claude Code", "오케스트레이션"), ("Planner", "설계"), ("Generator", "구현"), ("Evaluator", "검증")]),
        ("데이터와 배포", 0.9, "#fef3c7", "#92400e",
         [("Node API", "프록시"), ("Gemini 2.5", "이미지 분석"), ("Supabase", "Auth/DB"), ("Vercel", "배포")]),
    ]
    for title, y, bg, accent, items in layers:
        rbox(ax, 0.4, y, 13.2, 1.85, "", bg, edge="#cbd5e1", tc="#0f172a", lw=1.5)
        ax.text(0.85, y + 1.45, title,
                fontsize=13, fontweight="bold", color="#334155", ha="left")
        for i, (name, desc) in enumerate(items):
            ix = 1.0 + i * 3.15
            rbox(ax, ix, y + 0.28, 2.55, 0.95, f"{name}\n{desc}", accent,
                 tc="#ffffff", fs=11)

    TECH_PATH.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(TECH_PATH, dpi=220, bbox_inches="tight",
                facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)
    return TECH_PATH


# ─────────────────────────────────────────────────────────────────────────────
# Diagram 3: Screen list
# ─────────────────────────────────────────────────────────────────────────────

def make_screen_list() -> Path:
    setup_plot()
    screens = [
        ("01", "온보딩",   "나이/성별/키/체중\n활동량과 목표 설정",    "#2563eb"),
        ("02", "인증",     "이메일 회원가입\n게스트 모드",             "#0284c7"),
        ("03", "홈",       "오늘 칼로리 현황\n식사 추가 FAB",          "#16a34a"),
        ("04", "분석/수정", "Gemini 분석 결과\n사용자 보정 후 저장",   "#22c55e"),
        ("05", "기록",     "날짜별 식단\n수정/삭제/재분석",            "#f59e0b"),
        ("06", "수동 입력", "AI 실패 시\n직접 식사 입력",              "#ef4444"),
        ("07", "재분석",   "저장 이미지 또는\n새 사진으로 재분석",     "#8b5cf6"),
        ("08", "통계",     "주간 평균\n목표 추세 차트",                "#0f766e"),
        ("09", "설정",     "목표/모델/프로필\n재설정",                 "#64748b"),
    ]
    fig, axes = plt.subplots(3, 3, figsize=(15, 10.5))
    fig.patch.set_facecolor("#f8fafc")
    fig.suptitle("Cal AI 화면 및 기능 목록",
                 fontsize=22, fontweight="bold", color="#0f172a", y=0.98)
    plt.subplots_adjust(hspace=0.25, wspace=0.2)

    for ax, (num, name, desc, color) in zip(axes.flatten(), screens):
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.axis("off")
        # Card background
        card = FancyBboxPatch((0.4, 0.5), 9.2, 9.0,
                              boxstyle="round,pad=0.05", linewidth=2,
                              edgecolor=color, facecolor="#ffffff", zorder=2)
        ax.add_patch(card)
        # Top color band
        band = FancyBboxPatch((0.4, 8.0), 9.2, 1.5,
                              boxstyle="round,pad=0.05", linewidth=0,
                              edgecolor=color, facecolor=color, zorder=3)
        ax.add_patch(band)
        # Number circle
        circ = plt.Circle((1.55, 8.75), 0.6, color="#ffffff", zorder=4)
        ax.add_patch(circ)
        ax.text(1.55, 8.75, num,
                ha="center", va="center", color=color,
                fontweight="bold", fontsize=11, zorder=5)
        # Screen name
        ax.text(5.0, 8.75, name,
                ha="center", va="center",
                fontsize=13.5, fontweight="bold", color="#ffffff", zorder=5)
        # Description
        ax.text(5.0, 4.2, desc,
                ha="center", va="center",
                fontsize=12, color="#334155", linespacing=1.6, zorder=4)

    SCREEN_PATH.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(SCREEN_PATH, dpi=220, bbox_inches="tight",
                facecolor=fig.get_facecolor(), pad_inches=0.15)
    plt.close(fig)
    return SCREEN_PATH


# ─────────────────────────────────────────────────────────────────────────────
# Word document helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color="0f172a"):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13, True, color)
    return p


def add_para(doc, text, size=10.5, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    set_font(run, size)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, 10.5)


def add_prompt_box(doc, title, body):
    p = doc.add_paragraph()
    r = p.add_run(f"▶ {title}")
    set_font(r, 10.5, True, "0f766e")
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Cm(0.4)
    box.paragraph_format.right_indent = Cm(0.2)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "f1f5f9")
    box._p.get_or_add_pPr().append(shd)
    run = box.add_run(body)
    set_font(run, 9)


def set_word_font(cell_or_doc, size):
    """Set Korean font on all runs in a table cell."""
    for p in cell_or_doc.paragraphs:
        for run in p.runs:
            set_font(run, size)


# ─────────────────────────────────────────────────────────────────────────────
# Build Word document
# ─────────────────────────────────────────────────────────────────────────────

def create_doc(arch: Path, tech: Path, screen: Path) -> Path:
    doc = Document()

    # ── Global style ─────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    # ── Cover ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("AI Agent 중간점검 과제물")
    set_font(r, 24, True, "0f172a")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run("Cal AI: 음식 사진 한 장으로 기록하는 AI 칼로리 트래커")
    set_font(r, 13, False, "16a34a")

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (k, v) in enumerate([
        ("학생", "이현용"),
        ("날짜", dt.date.today().strftime("%Y년 %m월 %d일")),
        ("과목", "빅데이터와 인공지능"),
        ("과정", "인천대 MBA"),
    ]):
        info.cell(row, 0).text = k
        info.cell(row, 1).text = v
        shade_cell(info.cell(row, 0), "e2e8f0")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. 소개
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. 소개", 1)

    # 1-1. 적용 업무와 선택 이유
    add_heading(doc, "1-1. 적용 업무와 선택 이유", 2, "166534")

    add_para(doc,
        "이번 과제에서 AI Agent를 적용한 업무는 '음식 사진 기반 칼로리 기록 앱' 개발이다. "
        "다이어트나 건강 관리를 시작한 사용자는 식단 기록이 중요하다는 것을 알지만, 실제로는 음식명을 검색하고 "
        "양을 추정해 매번 입력하는 과정이 번거로워 쉽게 중단한다. 그래서 사용자가 사진을 찍으면 Gemini가 음식명과 "
        "칼로리, 단백질, 탄수화물, 지방을 추정하고, 사용자는 결과를 한 화면에서 수정한 뒤 저장하는 흐름을 만들었다."
    )

    # ★ 추가: 개인 다이어트 관련 문단 (MBA 업무와 무관하지만 개인 목적)
    add_para(doc,
        "다만, 이번 과제에서 선택한 주제는 현재 내가 수행하는 실제 업무와 직접적인 연관이 있는 주제는 아니다. "
        "개인적으로 진행 중인 다이어트와 건강 관리 습관을 개선하는 데 도움을 받기 위해 이 주제를 선택했다. "
        "매일 식단을 기록하고 칼로리를 확인하는 과정에서 꼭 필요한 기능만 자동화하면, "
        "MBA 수업에서 배운 AI Agent 접근 방법을 개인 생활의 문제 해결에도 충분히 적용할 수 있다고 판단했다."
    )
    # ★ 삭제된 문단: "이 주제를 선택한 이유는 MBA 수업에서 배운 AI Agent의 가치가..."
    #    → 해당 내용은 의도적으로 제외

    # 1-2. 기존 문제와 기대 효과
    add_heading(doc, "1-2. 기존 문제와 기대 효과", 2, "166534")

    # ★ 기존 테이블 삭제 → 새 내용으로 교체
    add_para(doc,
        "이 주제를 선택한 이유는 기존 칼로리 계산기 앱에는 내가 실제로 사용하지 않는 기능이 지나치게 많이 포함되어 있고, "
        "월 사용 비용도 약 4,000원에서 12,000원 수준으로 부담이 크기 때문이다. 그래서 실제로 내가 필요한 "
        "칼로리 계산 기능과 식단 기록 기능만 구현하고, 기타 불필요한 기능은 제외한 자동화 앱으로 직접 설계했다. "
        "실제 사용 비용도 이미지 분석 기준으로 한 달에 약 3장, 약 400원 정도가 소요될 것으로 예상한다."
    )

    # 1-3. 전체 아키텍처 개요
    add_heading(doc, "1-3. 전체 아키텍처 개요", 2, "166534")
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(arch), width=Inches(6.3))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. 진행 방법
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. 진행 방법: 어떻게 만들었나", 1)

    add_heading(doc, "2-1. 만든 스킬과 에이전트 목록", 2, "0f766e")
    add_bullets(doc, [
        "CLAUDE.md: 전체 파이프라인을 조율하는 마스터 오케스트레이션 지침",
        "agents/planner.md: PRD를 읽고 화면, 데이터, 기능 요구사항을 SPEC.md로 구조화",
        "agents/generator.md: SPEC과 QA 피드백을 바탕으로 output/index.html 단일 SPA 구현",
        "agents/evaluator.md: evaluation_criteria.md 기준으로 기능성, UX, 기술 안정성, 완성도 점검",
        "agents/evaluation_criteria.md: Generator와 Evaluator가 공유하는 채점 기준 문서",
    ])

    add_heading(doc, "2-2. 기술 스택", 2, "0f766e")
    tech_p = doc.add_paragraph()
    tech_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tech_p.add_run().add_picture(str(tech), width=Inches(6.3))
    add_para(doc,
        "프론트엔드는 HTML/CSS/Vanilla JS 기반의 모바일 우선 SPA로 구성했다. AI 분석은 Gemini 2.5 Flash와 Pro를 "
        "서버 API를 통해 호출하고, 인증과 식단 기록은 Supabase Auth와 PostgreSQL을 사용했다. 배포는 Vercel을 사용했으며 "
        "API 키는 클라이언트에 노출하지 않고 서버 환경변수로 관리했다."
    )

    add_heading(doc, "2-3. 사용한 주요 프롬프트", 2, "0f766e")
    add_prompt_box(doc, "Planner 호출",
        "agents/planner.md와 evaluation_criteria.md를 읽고, PRD.md의 요구사항을 기반으로 "
        "칼로리 트래커 상세 화면 설계서를 SPEC.md로 작성하라.")
    add_prompt_box(doc, "Generator 호출",
        "agents/generator.md를 읽고 SPEC.md 기반으로 Cal AI 전체를 output/index.html에 구현하라. "
        "Supabase JS CDN, Gemini API 호출, 모바일 Bottom Navigation을 포함하라.")
    add_prompt_box(doc, "Evaluator 호출",
        "agents/evaluator.md와 evaluation_criteria.md를 읽고 SPEC.md와 output/index.html을 비교 검증하라. "
        "기능성, UX, 기술 안정성, 완성도 점수를 QA_REPORT.md로 작성하라.")
    add_prompt_box(doc, "배포 환경 수정",
        "Vercel에서 /api/config와 /api/analyze가 정상 인식되도록 vercel.json과 API 런타임 설정을 수정하고, "
        "Gemini 키는 서버 .env로 관리하라.")

    add_heading(doc, "2-4. 수정 과정", 2, "0f766e")
    revisions = [
        ("R1", "초기 MVP 생성",      "Auth, 홈, 기록, 설정 중심의 기본 SPA 생성. QA 8.7/10."),
        ("R2", "UX 피드백 반영",     "회원가입 성공 메시지와 식사 카드 썸네일 개선. QA 9.2/10."),
        ("R3", "PRD v2 확장",        "온보딩, TDEE, 재분석, 수동 입력 fallback, 통계 화면 추가. QA 8.1/10."),
        ("R4", "서버 프록시 전환",   "Gemini API 키를 서버 .env로 이동, Supabase RLS 테이블 구성. QA 9.3/10."),
        ("R5", "UX 상용화 및 배포",  "카메라/갤러리 바텀시트, 커스텀 모달, Vercel 배포 완료. QA 9.6/10."),
    ]
    rev = doc.add_table(rows=1 + len(revisions), cols=3)
    rev.style = "Table Grid"
    for i, h in enumerate(["차수", "작업", "결과"]):
        rev.cell(0, i).text = h
        shade_cell(rev.cell(0, i), "0f766e")
        for run in rev.cell(0, i).paragraphs[0].runs:
            set_font(run, 10.5, True, "ffffff")
    for r, row in enumerate(revisions, 1):
        for c, val in enumerate(row):
            rev.cell(r, c).text = val
            set_word_font(rev.cell(r, c), 10)

    add_heading(doc, "2-5. 화면 목록", 2, "0f766e")
    screen_p = doc.add_paragraph()
    screen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screen_p.add_run().add_picture(str(screen), width=Inches(6.3))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. 결과 및 배운점
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. 결과 및 배운점", 1)

    add_heading(doc, "3-1. Before / After", 2, "7c3aed")
    ba = doc.add_table(rows=2, cols=2)
    ba.style = "Table Grid"
    for col, (label, color) in enumerate([("Before", "1e3a5f"), ("After", "166534")]):
        ba.cell(0, col).text = label
        shade_cell(ba.cell(0, col), color)
        for run in ba.cell(0, col).paragraphs[0].runs:
            set_font(run, 10.5, True, "ffffff")
    ba.cell(1, 0).text = (
        "식단 기록은 검색과 수동 입력 중심이었다. "
        "개발 과정도 아이디어, 구현, 검증이 섞여 있어 품질 개선 포인트를 분리하기 어려웠다."
    )
    ba.cell(1, 1).text = (
        "음식 사진 분석, 수동 보정, Supabase 저장, 재분석, 목표 설정, Vercel 배포까지 가능한 MVP가 완성되었다. "
        "최종 QA 점수는 9.6/10이다."
    )
    for c in (0, 1):
        set_word_font(ba.cell(1, c), 10)

    add_heading(doc, "3-2. 배운점과 꿀팁", 2, "7c3aed")
    add_bullets(doc, [
        "만드는 AI와 평가하는 AI를 분리하면 품질 관리가 훨씬 쉬워진다.",
        "좋은 PRD와 SPEC은 프롬프트보다 중요하다. 모호한 요구사항은 코드 품질의 변동으로 바로 이어진다.",
        "평가 기준을 Generator와 Evaluator가 함께 읽게 하면 목표가 일치한다.",
        "CHANGELOG와 QA_REPORT를 남기면 다음 반복에서 맥락 복원이 빠르다.",
        "API 키와 민감 정보는 초기에 서버 환경변수로 분리하는 편이 안전하다.",
    ])

    add_heading(doc, "3-3. 아직 막힌 부분 (도움 필요)", 2, "7c3aed")
    add_bullets(doc, [
        "현재 이미지는 단기적으로 base64 형태를 활용하므로, 장기 운영에는 Supabase Storage 전환이 필요하다.",
        "통계 화면은 기본 집계 중심이라 주간/월간 분석과 개인화 인사이트는 Phase 2로 남아 있다.",
        "Vercel 런타임 설정은 여러 번 시행착오가 있었고, 배포 환경별 최소 설정을 별도 문서화할 필요가 있다.",
    ])

    add_heading(doc, "3-4. 실패 경험 (제일 값진 내용)", 2, "7c3aed")
    add_para(doc,
        "가장 값진 실패는 R3에서 image_url을 null로 저장한 문제였다. 분석 단계에는 이미지 base64가 존재했지만 저장 단계에서 "
        "이를 meal_logs에 연결하지 않아 '저장된 이미지로 재분석' 기능이 비활성화되었다. 코드상으로는 앱이 정상 동작하는 것처럼 "
        "보였지만, 실제 사용자 시나리오에서는 핵심 기능 하나가 빠진 셈이었다. 이후 Evaluator가 단순 정적 점검뿐 아니라 데이터 "
        "흐름 관점에서 저장값과 후속 기능을 함께 검증해야 한다는 교훈을 얻었다."
    )
    add_para(doc,
        "또 다른 실패는 Vercel 배포 설정이다. 처음에는 빌드 설정과 서버리스 함수 인식 방식이 맞지 않아 /api 경로가 정상 동작하지 "
        "않았다. 최종적으로는 API 파일 구조와 런타임 설정을 단순화해 해결했다. 배포 문제는 코드가 아니라 환경의 계약 문제이므로, "
        "작동하는 최소 구성부터 확인해야 한다는 점을 배웠다."
    )

    add_heading(doc, "3-5. 앞으로의 적용 계획", 2, "7c3aed")
    add_bullets(doc, [
        "Phase 2에서 Supabase Storage, 주간/월간 통계, 연속 기록 스트릭을 추가한다.",
        "Playwright 기반 실제 브라우저 테스트를 Evaluator 루프에 포함해 정적 분석의 한계를 보완한다.",
        "같은 Planner-Generator-Evaluator 구조를 주간 보고서 생성, 회의록 자동화, 데이터 분석 리포트 작성 업무에도 적용한다.",
        "MBA 동기들에게 배포 URL을 공유해 실제 사용 피드백을 수집하고 다음 반복에 반영한다.",
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # 부록. 배포 정보와 체크리스트
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "부록 A. 배포 정보와 체크리스트", 1)
    add_para(doc, "배포 URL: https://calorie-ai-gamma.vercel.app")
    add_bullets(doc, [
        "소개, 진행 방법, 결과 및 배운점 섹션 작성 완료",
        "아키텍처, 기술 스택, 화면 목록 이미지 포함",
        "Before/After와 실패 경험 포함",
        "API 키 등 민감 정보는 문서에 포함하지 않음",
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # 부록. 실제 화면 캡처
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    app_h = add_heading(doc, "부록 B. 실제 화면 캡처", 1)
    add_para(doc,
        "아래 이미지는 Cal AI 앱의 실제 사용 흐름에 따라 로그인, 프로필 설정, 주요 기능 화면 순서로 정리했으며, "
        "한 페이지에 두 장씩 배치하였다."
    )

    def remove_table_borders(tbl):
        tbl_pr = tbl._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = f"w:{edge}"
            el = borders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                borders.append(el)
            el.set(qn("w:val"), "nil")

    def photo_width(path: Path) -> Inches:
        with Image.open(path) as img:
            w, h = img.size
        return Inches(2.72) if h > w else Inches(5.5)

    for pair_start in range(0, len(PHOTO_ORDER), 2):
        if pair_start > 0:
            doc.add_page_break()

        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        remove_table_borders(tbl)

        for col in range(2):
            item_idx = pair_start + col
            if item_idx >= len(PHOTO_ORDER):
                break
            caption, filename = PHOTO_ORDER[item_idx]
            cell = tbl.cell(0, col)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.width = Inches(3.1)
            cell.text = ""

            # Caption
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(4)
            cr = cp.add_run(f"부록 B-{item_idx + 1}. {caption}")
            set_font(cr, 9.5, True)

            # Image
            img_path = PHOTO_DIR / filename
            ip = cell.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.add_run().add_picture(str(img_path), width=photo_width(img_path))

            # Filename label
            fnp = cell.add_paragraph()
            fnp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fnr = fnp.add_run(filename)
            set_font(fnr, 7.5)

    # Save
    OUTPUT_STANDARD.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_STANDARD)
    shutil.copyfile(OUTPUT_STANDARD, OUTPUT_DOCX)
    return OUTPUT_DOCX


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("[1] 아키텍처 다이어그램 생성 중...")
    arch = make_architecture()
    print(f"   -> {arch}")

    print("[2] 기술 스택 다이어그램 생성 중...")
    tech = make_tech_stack()
    print(f"   -> {tech}")

    print("[3] 화면 목록 다이어그램 생성 중...")
    screen = make_screen_list()
    print(f"   -> {screen}")

    print("[4] Word 문서 생성 중...")
    docx = create_doc(arch, tech, screen)
    print(f"\n완료: {docx}")
