# -*- coding: utf-8 -*-
"""Update the appendix-included MBA report text and diagrams."""
from __future__ import annotations

import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parent
SOURCE_DOC = BASE_DIR / "MBA_중간점검_과제물_이현용_부록포함.standard-docx"
OUTPUT_STANDARD = SOURCE_DOC
OUTPUT_DOCX = BASE_DIR / "MBA_중간점검_과제물_이현용_부록포함_수정본.docx"

ARCH_PATH = BASE_DIR / "architecture_diagram.png"
TECH_PATH = BASE_DIR / "tech_stack_diagram.png"
SCREEN_PATH = BASE_DIR / "screen_list_diagram.png"


def setup_plot() -> None:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 180
    plt.rcParams["savefig.dpi"] = 260


def rounded_box(ax, x, y, w, h, text, face, edge=None, text_color="#0f172a", size=13, radius=0.08):
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0.04,rounding_size={radius}",
        linewidth=2,
        edgecolor=edge or face,
        facecolor=face,
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        color=text_color,
        fontsize=size,
        fontweight="bold",
        linespacing=1.28,
    )


def arrow(ax, start, end, color="#2563eb", width=2.5, rad=0.0):
    ax.annotate(
        "",
        xy=end,
        xytext=start,
        arrowprops={
            "arrowstyle": "->",
            "lw": width,
            "color": color,
            "shrinkA": 5,
            "shrinkB": 5,
            "connectionstyle": f"arc3,rad={rad}",
        },
    )


def make_architecture() -> None:
    setup_plot()
    fig, ax = plt.subplots(figsize=(15.5, 8.7))
    fig.patch.set_facecolor("#f8fafc")
    ax.set_xlim(0, 15.5)
    ax.set_ylim(0, 8.7)
    ax.axis("off")

    ax.text(
        7.75,
        8.25,
        "Cal AI 3-Agent 파이프라인 아키텍처",
        ha="center",
        fontsize=24,
        fontweight="bold",
        color="#0f172a",
    )

    y_main = 4.25
    rounded_box(ax, 0.45, y_main, 1.9, 1.05, "사용자\n요청", "#475569", text_color="#ffffff")
    rounded_box(ax, 3.0, y_main, 2.2, 1.05, "Planner\nAgent", "#0284c7", text_color="#ffffff")
    rounded_box(ax, 6.15, y_main, 2.2, 1.05, "Generator\nAgent", "#16a34a", text_color="#ffffff")
    rounded_box(ax, 9.35, y_main, 2.2, 1.05, "Evaluator\nAgent", "#f59e0b", text_color="#ffffff")
    rounded_box(ax, 12.65, y_main, 2.2, 1.05, "합격\n완료", "#22c55e", text_color="#ffffff")

    rounded_box(ax, 3.1, 6.05, 2.0, 0.85, "PRD.md", "#e0f2fe", edge="#0284c7", text_color="#075985", size=12)
    rounded_box(ax, 6.25, 6.05, 2.0, 0.85, "SPEC.md", "#dcfce7", edge="#16a34a", text_color="#166534", size=12)
    rounded_box(ax, 6.25, 2.35, 2.0, 0.85, "output/\nindex.html", "#f8fafc", edge="#64748b", text_color="#334155", size=11)
    rounded_box(ax, 9.25, 2.35, 2.35, 0.85, "QA_REPORT.md\n점수/피드백", "#ffedd5", edge="#f59e0b", text_color="#9a3412", size=10)

    arrow(ax, (2.35, 4.78), (3.0, 4.78))
    arrow(ax, (5.2, 4.78), (6.15, 4.78))
    arrow(ax, (8.35, 4.78), (9.35, 4.78))
    arrow(ax, (11.55, 4.78), (12.65, 4.78))

    arrow(ax, (4.1, 6.05), (6.18, 5.2), "#64748b", width=2.2)
    arrow(ax, (7.25, 6.05), (7.25, 5.3), "#64748b", width=2.2)
    arrow(ax, (10.43, 4.25), (10.43, 3.2), "#64748b", width=2.2)
    arrow(ax, (9.25, 2.78), (8.25, 2.78), "#ef4444", width=2.2)
    arrow(ax, (7.25, 3.2), (7.25, 4.25), "#ef4444", width=2.2)
    ax.text(
        7.75,
        1.98,
        "불합격 또는 조건부 합격 시 피드백 반영",
        ha="center",
        fontsize=12,
        color="#ef4444",
        fontweight="bold",
    )

    ax.plot([0.65, 14.85], [1.55, 1.55], linestyle=(0, (5, 5)), color="#cbd5e1", linewidth=2)
    stack = [
        ("HTML/CSS/JS\n모바일 SPA", "#0ea5e9"),
        ("Gemini 2.5\n이미지 분석", "#7c3aed"),
        ("Supabase\nAuth + DB", "#059669"),
        ("Vercel\n배포 + API", "#111827"),
    ]
    for i, (label, color) in enumerate(stack):
        rounded_box(ax, 1.0 + i * 3.55, 0.35, 2.65, 0.85, label, color, text_color="#ffffff", size=12)

    fig.savefig(ARCH_PATH, bbox_inches="tight", facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)


def make_tech_stack() -> None:
    setup_plot()
    fig, ax = plt.subplots(figsize=(13.5, 8.2))
    fig.patch.set_facecolor("#ffffff")
    ax.set_xlim(0, 13.5)
    ax.set_ylim(0, 8.2)
    ax.axis("off")
    ax.text(6.75, 7.78, "Cal AI 기술 스택", ha="center", fontsize=23, fontweight="bold", color="#0f172a")

    layers = [
        ("사용자 화면", 5.45, "#dbeafe", "#1d4ed8", [("HTML5", "구조"), ("CSS3", "모바일 UI"), ("Vanilla JS", "SPA 상태"), ("SVG/Canvas", "차트")]),
        ("AI Agent 개발", 3.15, "#dcfce7", "#15803d", [("Claude Code", "오케스트레이션"), ("Planner", "설계"), ("Generator", "구현"), ("Evaluator", "검증")]),
        ("데이터와 배포", 0.85, "#fef3c7", "#92400e", [("Node API", "프록시"), ("Gemini", "분석"), ("Supabase", "Auth/DB"), ("Vercel", "배포")]),
    ]
    for title, y, bg, accent, items in layers:
        rounded_box(ax, 0.45, y, 12.6, 1.75, "", bg, edge="#cbd5e1")
        ax.text(0.85, y + 1.35, title, fontsize=14, fontweight="bold", color="#334155", ha="left")
        for i, (name, desc) in enumerate(items):
            x = 1.0 + i * 3.0
            rounded_box(ax, x, y + 0.25, 2.35, 0.82, f"{name}\n{desc}", accent, text_color="#ffffff", size=11)

    fig.savefig(TECH_PATH, bbox_inches="tight", facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)


def make_screen_list() -> None:
    setup_plot()
    screens = [
        ("01", "온보딩", "나이/성별/키/체중\n활동량과 목표 설정", "#2563eb"),
        ("02", "인증", "이메일 회원가입\n게스트 모드", "#0284c7"),
        ("03", "홈", "오늘 칼로리 현황\n식사 추가 FAB", "#16a34a"),
        ("04", "분석/수정", "Gemini 분석 결과\n사용자 보정 후 저장", "#22c55e"),
        ("05", "기록", "날짜별 식단\n수정/삭제/재분석", "#f59e0b"),
        ("06", "수동 입력", "AI 실패 시\n직접 식사 입력", "#ef4444"),
        ("07", "재분석", "저장 이미지 또는\n새 사진으로 재분석", "#8b5cf6"),
        ("08", "통계", "주간 평균\n목표 추세 차트", "#0f766e"),
        ("09", "설정", "목표/모델/프로필\n재설정", "#64748b"),
    ]
    fig, axes = plt.subplots(3, 3, figsize=(14, 9.6))
    fig.patch.set_facecolor("#f8fafc")
    fig.suptitle("Cal AI 화면 및 기능 목록", fontsize=23, fontweight="bold", color="#0f172a", y=0.98)
    for ax, (num, name, desc, color) in zip(axes.flatten(), screens):
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.axis("off")
        rounded_box(ax, 0.5, 0.6, 9.0, 8.8, "", "#ffffff", edge="#d8e0ea")
        circ = plt.Circle((1.65, 8.05), 0.8, color=color)
        ax.add_patch(circ)
        ax.text(1.65, 8.05, num, ha="center", va="center", color="#ffffff", fontweight="bold", fontsize=11)
        ax.text(2.75, 8.05, name, ha="left", va="center", fontsize=14, fontweight="bold", color="#0f172a")
        ax.plot([1.05, 8.95], [6.55, 6.55], color=color, linewidth=2, alpha=0.32)
        ax.text(5.0, 4.0, desc, ha="center", va="center", fontsize=12, color="#334155", linespacing=1.55)

    fig.savefig(SCREEN_PATH, bbox_inches="tight", facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)


def set_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run, 10.5)


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    replace_paragraph_text(new_paragraph, text)
    return new_paragraph


def remove_existing_problem_table(doc: Document) -> None:
    for table in list(doc.tables):
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Before: 기존 문제" in table_text or "선택 이유와 기대 효과" in table_text:
            table._element.getparent().remove(table._element)
            return


def replace_picture_after_heading(doc: Document, heading: str, image_path: Path, width: float) -> None:
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == heading)
    for paragraph in doc.paragraphs[start + 1 :]:
        if paragraph.text.strip().startswith(("1-", "2-", "3.", "부록.")):
            break
        if paragraph._element.xpath(".//w:drawing"):
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(width))
            return
    raise ValueError(f"Image paragraph after heading not found: {heading}")


def update_report() -> None:
    make_architecture()
    make_tech_stack()
    make_screen_list()

    doc = Document(SOURCE_DOC)

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "1-1. 적용 업무와 선택 이유":
            replace_paragraph_text(
                doc.paragraphs[i + 2],
                "또한 이 주제는 현재 내가 수행하는 실제 업무와 직접적인 관련이 있는 과제는 아니지만, "
                "개인적으로 진행 중인 다이어트와 건강 관리 습관을 개선하는 데 도움을 받기 위해 선택했다. "
                "매일 식단을 기록하고 칼로리를 확인하는 과정에서 꼭 필요한 기능만 자동화하면, "
                "업무 자동화에서 배운 AI Agent 접근법을 개인 생활의 문제 해결에도 적용할 수 있다고 판단했다.",
            )
        if text == "1-2. 기존 문제와 기대 효과":
            remove_existing_problem_table(doc)
            already_added = doc.paragraphs[i + 1].text.strip().startswith("이 주제를 선택한 이유는 기존 칼로리 계산기 앱")
            if not already_added:
                insert_paragraph_after(
                    paragraph,
                    "이 주제를 선택한 이유는 기존 칼로리 계산기 앱에는 내가 사용하지 않는 기능이 지나치게 많이 포함되어 있고, "
                    "월 사용 비용도 약 4,000원에서 12,000원 수준으로 부담이 크기 때문이다. 그래서 실제로 내가 필요한 "
                    "칼로리 계산 기능과 식단 기록 기능만 추가하고, 기타 불필요한 기능은 구현하지 않은 자동화 앱으로 설계했다. "
                    "실제 사용 비용도 이미지 분석 기준으로 한 달에 약 3장, 약 400원 정도가 소요될 것으로 예상한다.",
                )

    replace_picture_after_heading(doc, "1-3. 전체 아키텍처 개요", ARCH_PATH, 6.3)
    replace_picture_after_heading(doc, "2-2. 기술 스택", TECH_PATH, 6.3)
    replace_picture_after_heading(doc, "2-5. 화면 목록", SCREEN_PATH, 6.3)

    doc.save(OUTPUT_STANDARD)
    shutil.copyfile(OUTPUT_STANDARD, OUTPUT_DOCX)


if __name__ == "__main__":
    update_report()
    print(OUTPUT_DOCX)
