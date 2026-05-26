# -*- coding: utf-8 -*-
"""MBA 중간점검 과제물 생성 스크립트"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
plt.rcParams['font.family'] = 'Malgun Gothic'
plt.rcParams['axes.unicode_minus'] = False
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────
# 1. 아키텍처 다이어그램
# ─────────────────────────────────────────────
def make_architecture():
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('#f0f7ff')
    ax.set_facecolor('#f0f7ff')

    def box(x, y, w, h, text, color, fontsize=10, textcolor='white'):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center',
                fontsize=fontsize, color=textcolor, fontweight='bold',
                wrap=True, multialignment='center')

    def arrow(x1, y1, x2, y2):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#1e40af', lw=2))

    # 타이틀
    ax.text(6, 6.6, '3-Agent 하네스 파이프라인 아키텍처\nCal AI 칼로리 트래커',
            ha='center', va='center', fontsize=14, fontweight='bold', color='#1e3a5f')

    # 사용자
    box(0.3, 3.1, 1.8, 0.8, '사용자\n요청', '#6366f1')
    # Planner
    box(2.6, 3.1, 2.0, 0.8, '① Planner\nAgent', '#0ea5e9')
    # SPEC
    box(5.2, 4.0, 1.6, 0.6, 'SPEC.md', '#64748b', textcolor='white', fontsize=9)
    # Generator
    box(5.2, 3.1, 2.0, 0.8, '② Generator\nAgent', '#10b981')
    # output
    box(5.2, 2.1, 1.6, 0.6, 'output/\nindex.html', '#64748b', textcolor='white', fontsize=9)
    # Evaluator
    box(7.8, 3.1, 2.0, 0.8, '③ Evaluator\nAgent', '#f59e0b')
    # QA Report
    box(7.8, 2.1, 2.0, 0.6, 'QA_REPORT.md\n판정(합격/불합격)', '#64748b', textcolor='white', fontsize=8)
    # 판정
    box(10.3, 3.1, 1.4, 0.8, '합격\n→ 완료', '#22c55e')

    # 화살표
    arrow(2.1, 3.5, 2.6, 3.5)
    arrow(4.6, 3.5, 5.2, 3.5)
    arrow(6.0, 4.0, 6.0, 3.9)   # spec → generator
    arrow(7.2, 3.5, 7.8, 3.5)
    arrow(8.8, 3.5, 10.3, 3.5)
    arrow(8.8, 2.4, 5.3, 2.4)   # 피드백 루프
    ax.annotate('', xy=(5.2, 3.1), xytext=(5.2, 2.4),
                arrowprops=dict(arrowstyle='->', color='#ef4444', lw=2))

    # 피드백 레이블
    ax.text(7.0, 2.1, '불합격 시 피드백 반영 (최대 3회)', ha='center', fontsize=8, color='#ef4444', style='italic')

    # 하단: 기술 스택 레이어
    box(0.3, 0.6, 2.5, 0.9, 'Gemini API\n(Flash / Pro)', '#7c3aed')
    box(3.2, 0.6, 2.5, 0.9, 'Supabase\n(Auth + DB)', '#0284c7')
    box(6.1, 0.6, 2.5, 0.9, 'Vercel\n(Edge Runtime)', '#374151')
    box(9.0, 0.6, 2.7, 0.9, 'Claude Code\n(AI Orchestrator)', '#0f172a')

    ax.text(6, 0.3, '기술 스택', ha='center', fontsize=9, color='#475569')

    # 점선 구분선
    ax.plot([0.3, 11.7], [1.65, 1.65], '--', color='#94a3b8', lw=1)

    plt.tight_layout(pad=0.5)
    path = os.path.join(OUT_DIR, 'architecture_diagram.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='#f0f7ff')
    plt.close()
    print('[OK] architecture_diagram.png 저장')
    return path


# ─────────────────────────────────────────────
# 2. 기술 스택 다이어그램
# ─────────────────────────────────────────────
def make_tech_stack():
    fig, ax = plt.subplots(figsize=(11, 8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 8)
    ax.axis('off')
    fig.patch.set_facecolor('#fafafa')

    ax.text(5.5, 7.6, '기술 스택 레이어드 다이어그램', ha='center', va='center',
            fontsize=14, fontweight='bold', color='#1e3a5f')

    layers = [
        {
            'label': '🖥  UI / 인터페이스 레이어',
            'y': 5.9, 'color': '#dbeafe',
            'items': [
                ('HTML5 + CSS3', '#3b82f6', '마크업 & 스타일'),
                ('Vanilla JS (SPA)', '#2563eb', '라우팅·이벤트'),
                ('SVG 링/바 차트', '#1d4ed8', '칼로리 시각화'),
                ('PWA (sw.js)', '#1e40af', '오프라인·설치'),
            ]
        },
        {
            'label': '🤖  AI / 에이전트 레이어',
            'y': 3.85, 'color': '#dcfce7',
            'items': [
                ('Gemini 2.5 Flash', '#16a34a', '기본 이미지 분석'),
                ('Gemini 2.5 Pro', '#15803d', '고정밀 분석 (선택)'),
                ('Claude Code', '#166534', '하네스 오케스트레이터'),
                ('Planner/Gen/Eval', '#14532d', '3-Agent 파이프라인'),
            ]
        },
        {
            'label': '🗄  데이터 / 서버 레이어',
            'y': 1.8, 'color': '#fef9c3',
            'items': [
                ('Supabase Auth', '#ca8a04', '이메일/소셜 인증'),
                ('Supabase DB (RLS)', '#a16207', 'meal_logs·profiles'),
                ('Node.js API Proxy', '#713f12', '/api/analyze·config'),
                ('Vercel Edge', '#92400e', '서버리스 배포'),
            ]
        },
    ]

    for layer in layers:
        # 레이어 배경
        rect = FancyBboxPatch((0.3, layer['y']), 10.4, 1.7,
                              boxstyle="round,pad=0.15",
                              facecolor=layer['color'], edgecolor='#cbd5e1', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(0.6, layer['y'] + 1.55, layer['label'],
                fontsize=10, fontweight='bold', color='#374151', va='top')

        for i, (name, color, desc) in enumerate(layer['items']):
            x = 0.7 + i * 2.55
            y = layer['y'] + 0.15
            item_rect = FancyBboxPatch((x, y), 2.2, 1.0,
                                       boxstyle="round,pad=0.08",
                                       facecolor=color, edgecolor='white', linewidth=1)
            ax.add_patch(item_rect)
            ax.text(x + 1.1, y + 0.65, name, ha='center', va='center',
                    fontsize=9, color='white', fontweight='bold')
            ax.text(x + 1.1, y + 0.25, desc, ha='center', va='center',
                    fontsize=7.5, color='#f0fdf4')

    plt.tight_layout(pad=0.5)
    path = os.path.join(OUT_DIR, 'tech_stack_diagram.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='#fafafa')
    plt.close()
    print('[OK] tech_stack_diagram.png 저장')
    return path


# ─────────────────────────────────────────────
# 3. 화면 목록 다이어그램
# ─────────────────────────────────────────────
def make_screen_list():
    screens = [
        ('01', '온보딩', '신체정보 입력\n(성별·나이·키·몸무게)\n→ TDEE 자동 계산', '#6366f1'),
        ('02', '인증', '이메일 회원가입/로그인\n게스트 모드 지원\nSupabase Auth', '#0ea5e9'),
        ('03', '홈 대시보드', '오늘 칼로리 링차트\n매크로 진행률\n음식 추가 FAB', '#10b981'),
        ('04', '음식 분석', '카메라/갤러리 촬영\nGemini AI 분석\n결과 수정 후 저장', '#22c55e'),
        ('05', '기록 히스토리', '날짜별 식단 리스트\n사진 썸네일 표시\n칼로리 상세 드롭다운', '#f59e0b'),
        ('06', '재분석', '저장 이미지 재전송\n또는 새 사진으로\n기존 기록 업데이트', '#ef4444'),
        ('07', '통계', '일/주/월 SVG 바차트\n목표 대비 추세선\n(Phase 2 준비 중)', '#8b5cf6'),
        ('08', '설정', 'TDEE 목표 수정\nGemini 모델 선택\n(Flash / Pro)', '#64748b'),
        ('09', '수동 입력 Fallback', 'AI 분석 실패 시\n음식명·영양소\n직접 입력 저장', '#94a3b8'),
    ]

    fig, axes = plt.subplots(3, 3, figsize=(13, 9))
    fig.patch.set_facecolor('#f8fafc')
    fig.suptitle('Cal AI 칼로리 트래커 — 화면 목록 (9개 화면)', fontsize=14, fontweight='bold', color='#1e3a5f', y=0.98)

    for idx, (ax, (num, name, desc, color)) in enumerate(zip(axes.flatten(), screens)):
        ax.set_facecolor(color + '18')
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.axis('off')

        # 번호 뱃지
        circle = plt.Circle((1.2, 8.2), 0.9, color=color, zorder=3)
        ax.add_patch(circle)
        ax.text(1.2, 8.2, num, ha='center', va='center', fontsize=11,
                fontweight='bold', color='white', zorder=4)

        # 화면명
        ax.text(2.8, 8.2, name, ha='left', va='center', fontsize=11,
                fontweight='bold', color='#1e293b')

        # 구분선
        ax.plot([0.5, 9.5], [6.8, 6.8], '-', color=color, lw=1.5, alpha=0.4)

        # 설명
        ax.text(5, 4.0, desc, ha='center', va='center', fontsize=9,
                color='#334155', multialignment='center',
                bbox=dict(boxstyle='round,pad=0.4', facecolor='white',
                          edgecolor=color, linewidth=1.2, alpha=0.9))

        # 테두리
        for spine in ['top', 'right', 'bottom', 'left']:
            ax.spines[spine].set_visible(False)
        rect = FancyBboxPatch((0.05, 0.05), 9.9, 9.9,
                              boxstyle="round,pad=0.1",
                              facecolor='none', edgecolor=color, linewidth=2)
        ax.add_patch(rect)

    plt.tight_layout(rect=[0, 0, 1, 0.97])
    path = os.path.join(OUT_DIR, 'screen_list_diagram.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='#f8fafc')
    plt.close()
    print('[OK] screen_list_diagram.png 저장')
    return path


# ─────────────────────────────────────────────
# 4. Word 문서 생성
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_heading_custom(doc, text, level, color_hex='1e3a5f'):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
    return p

def add_highlight_para(doc, text, bg='e8f5e9'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg)
    shd.set(qn('w:val'), 'clear')
    pPr = p._p.get_or_add_pPr()
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def create_word_doc(arch_path, tech_path, screen_path):
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)

    # ── 표지 ──
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('AI Agent 중간점검 과제물')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    run.font.name = 'Malgun Gothic'

    doc.add_paragraph()
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run('Cal AI — 음식 사진 한 장으로 칼로리를 기록하는 AI 칼로리 트래커')
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)
    run2.font.name = 'Malgun Gothic'

    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('학생', '이현용'),
        ('날짜', datetime.date.today().strftime('%Y년 %m월 %d일')),
        ('과목', '빅데이터와 인공지능'),
        ('과정', '인천대학교 MBA'),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v
        info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ── ① 소개 ──
    add_heading_custom(doc, '① 소개', 1)

    add_heading_custom(doc, '1-1. 적용 업무 선택 — 왜 칼로리 트래커인가?', 2, '166534')
    doc.add_paragraph(
        '다이어트를 시작한 직장인들이 식단 기록을 포기하는 가장 큰 이유는 "번거로움"입니다. '
        '기존 칼로리 앱은 음식명을 직접 검색하고, 양을 그램 단위로 입력해야 해서 하루 3끼를 기록하는 데 '
        '10~15분이 걸립니다. 저 역시 이 문제를 겪고 있었습니다.\n\n'
        'AI 에이전트로 이 문제를 해결할 수 있다는 가설을 세웠습니다. '
        '"음식 사진 한 장 → Gemini AI 자동 분석 → 10초 내 저장"이라는 흐름이 가능하다면, '
        '기록 습관 형성의 진입 장벽을 낮출 수 있습니다. '
        'MBA 수업에서 배운 AI 에이전트 개념을 실제 문제에 적용해 보고 싶었습니다.'
    )

    add_heading_custom(doc, '1-2. 기존 문제와 에이전트 적용 후 기대효과', 2, '166534')
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_bg(tbl.cell(0, 0), '1e3a5f')
    set_cell_bg(tbl.cell(0, 1), '166534')
    h0 = tbl.cell(0, 0).paragraphs[0]
    h0.add_run('Before — 기존 문제').font.color.rgb = RGBColor(255, 255, 255)
    h0.runs[0].font.bold = True
    h1 = tbl.cell(0, 1).paragraphs[0]
    h1.add_run('After — 에이전트 적용 기대효과').font.color.rgb = RGBColor(255, 255, 255)
    h1.runs[0].font.bold = True
    tbl.cell(1, 0).text = (
        '• 음식명 검색 + 양 입력에 10분 이상 소요\n'
        '• 기록 포기율 높음 (귀찮음)\n'
        '• 앱 개발 전문 지식 없어 직접 개발 불가\n'
        '• Planner-Generator-Evaluator 분리 없이\n  단일 프롬프트로 복잡한 앱 개발 시 일관성 부족'
    )
    tbl.cell(1, 1).text = (
        '• 사진 촬영 → 저장 10초 이내 완료\n'
        '• AI 분석 + 수동 수정 병행으로 신뢰도 확보\n'
        '• 3-Agent 하네스로 코딩 없이 SPA 자동 생성\n'
        '• Evaluator가 품질 보증 → 반복 개선 자동화\n'
        '• Vercel 배포로 누구나 접속 가능한 앱 완성'
    )

    doc.add_paragraph()
    add_heading_custom(doc, '1-3. 전체 아키텍처 개요', 2, '166534')
    doc.add_picture(arch_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('[ 그림 1 ] 3-Agent 하네스 파이프라인: Planner → Generator → Evaluator')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_page_break()

    # ── ② 진행 방법 ──
    add_heading_custom(doc, '② 진행 방법 — 어떻게 만들었나', 1)

    add_heading_custom(doc, '2-1. 만든 스킬(에이전트) 목록', 2, '0f766e')
    skills = [
        ('CLAUDE.md 하네스 오케스트레이터',
         '전체 파이프라인을 조율하는 마스터 설정 파일. '
         'Planner → Generator → Evaluator 순서를 정의하고, 각 단계에서 Task 도구로 서브에이전트를 호출합니다. '
         '합격 판정까지 최대 3회 반복 루프를 자동 실행합니다.'),
        ('agents/planner.md (Planner 에이전트)',
         'PRD.md를 읽고 Generator가 구현 가능한 상세 화면 설계서(SPEC.md)를 작성합니다. '
         '"무엇을 만들지"와 "어떻게 보일지"에 집중하며, 기술 구현 방법은 정하지 않습니다. '
         '5개 화면(인증/홈/히스토리/통계/설정) 구조와 디자인 방향을 정의합니다.'),
        ('agents/generator.md (Generator 에이전트)',
         'SPEC.md와 QA 피드백을 바탕으로 output/index.html(단일 파일 SPA)을 생성합니다. '
         'Supabase JS CDN, Gemini API 연동, 모바일 퍼스트 레이아웃을 포함하며 '
         'SELF_CHECK.md를 작성해 자체 검증합니다.'),
        ('agents/evaluator.md (Evaluator 에이전트)',
         'evaluation_criteria.md의 4개 항목(기능성 40%·UX 30%·기술품질 20%·완성도 10%)으로 '
         'output/index.html을 채점합니다. 합격(7.0 이상) / 조건부 / 불합격 판정 후 '
         '구체적인 개선 지시를 QA_REPORT.md로 출력합니다.'),
        ('agents/evaluation_criteria.md (품질 기준 문서)',
         '4개 채점 항목의 합격/불합격 기준을 정의한 공유 문서. '
         'Generator는 이 기준을 미리 읽고 구현하고, Evaluator는 동일 기준으로 채점합니다. '
         '"만드는 AI와 평가하는 AI를 분리"하는 핵심 구조입니다.'),
    ]
    for i, (name, desc) in enumerate(skills, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run_num = p.add_run(f'[스킬 {i}] ')
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(0x0e, 0xa5, 0xe9)
        run_name = p.add_run(name + '\n')
        run_name.font.bold = True
        p.add_run(desc)

    add_heading_custom(doc, '2-2. 기술 스택', 2, '0f766e')
    doc.add_picture(tech_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2 = doc.add_paragraph('[ 그림 2 ] 기술 스택 레이어드 다이어그램 (UI → AI → 데이터/서버)')
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2.runs[0].font.size = Pt(9)
    cap2.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    add_heading_custom(doc, '2-3. 사용한 주요 프롬프트', 2, '0f766e')
    prompts = [
        ('① 첫 생성 요청 (Planner 호출)',
         'agents/planner.md 파일을 읽고, 그 지시를 따라라.\n'
         'agents/evaluation_criteria.md 파일도 읽고 참고하라.\n'
         'PRD.md 파일을 읽어라. 이것이 제품 요구사항이다.\n'
         'PRD의 MVP 기능을 기반으로 상세 화면 설계서를 작성하라.\n'
         '결과를 SPEC.md 파일로 저장하라.'),
        ('② 초기 구현 요청 (Generator 호출)',
         'agents/generator.md 파일을 읽고, 그 지시를 따라라.\n'
         'SPEC.md 파일을 읽고, 칼로리 트래커 앱 전체를 한 번에 구현하라.\n'
         '반드시: Supabase JS CDN 포함 / Gemini API를 fetch로 호출 (base64 이미지 방식) /\n'
         'API 키는 localStorage에서 읽어라 (하드코딩 금지) / 모바일 퍼스트, Bottom Navigation 포함\n'
         '결과를 output/index.html 파일로 저장하라. 완료 후 SELF_CHECK.md를 작성하라.'),
        ('③ Evaluator 검수 요청',
         'agents/evaluator.md 파일을 읽고, 그 지시를 따라라.\n'
         'agents/evaluation_criteria.md 파일을 읽어라. 이것이 채점 기준이다.\n'
         'SPEC.md + output/index.html을 읽고 4개 항목을 채점하라.\n'
         '최종 판정(합격/조건부/불합격)을 내리고 QA_REPORT.md로 저장하라.'),
        ('④ PRD v2 방향 전환 요청',
         'amirdora/ai_calorie_tracker 저장소를 참고하여 PRD를 전면 재작성해줘.\n'
         '온보딩(Harris-Benedict TDEE 자동계산) + Gemini 이중 모델(Flash/Pro) +\n'
         '재분석 UX(저장 이미지 재전송 vs 새 사진)를 포함해줘.\n'
         'MVP를 6기능 → 9기능으로 확장하고 PRD.md로 저장해줘.'),
        ('⑤ 배포 환경 수정 요청',
         'Vercel에서 API 라우트가 인식되지 않는 문제를 해결해줘.\n'
         'vercel.json의 builds 배열 구조를 수정하고, api/analyze.js와 api/config.js에\n'
         'Edge Runtime 설정(export const config = { runtime: \'edge\' })을 추가해줘.\n'
         'Gemini API 키는 클라이언트에 노출되지 않도록 서버 .env로 관리해줘.'),
    ]
    for purpose, prompt_text in prompts:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(6)
        run_h = p_head.add_run(f'▶ {purpose}')
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

        p_box = doc.add_paragraph()
        p_box.paragraph_format.left_indent = Cm(0.8)
        p_box.paragraph_format.right_indent = Cm(0.5)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'f1f5f9')
        shd.set(qn('w:val'), 'clear')
        pPr = p_box._p.get_or_add_pPr()
        pPr.append(shd)
        run_p = p_box.add_run(prompt_text)
        run_p.font.size = Pt(9)
        run_p.font.name = 'Courier New'

    add_heading_custom(doc, '2-4. 수정 과정 (반복 개선 이력)', 2, '0f766e')
    revisions = [
        ('R1 (2026-04-20)',  'MVP 초기 생성',         '하네스 1회 실행. 기본 SPA 생성 (Auth+홈+히스토리+설정). QA 점수 8.7/10, 합격'),
        ('R2 (2026-04-20)',  'UX 개선 피드백 반영',   'Evaluator 피드백: 회원가입 성공 메시지 빨간색 → 초록 분리. 식사 카드 썸네일 추가. 9.2/10'),
        ('R3 (2026-04-21)',  'PRD v2 기반 전면 재구현','9개 화면, Harris-Benedict TDEE, 재분석 모달, 수동입력 Fallback 추가. 8.1/10'),
        ('R4 (2026-04-22)',  'Supabase 테이블 직접 생성','MCP로 profiles·meal_logs 테이블 생성. Gemini 모델별 API 키 분리(Flash/Pro)'),
        ('R5 (2026-04-23)',  '상업화 UX + Vercel 배포','바텀시트, 커스텀 모달 4종, FAB 카메라 아이콘, calorie-ai-gamma.vercel.app 배포 완료'),
    ]
    rev_tbl = doc.add_table(rows=len(revisions)+1, cols=3)
    rev_tbl.style = 'Table Grid'
    rev_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['회차 / 날짜', '작업 내용', '결과']
    for i, h in enumerate(headers):
        c = rev_tbl.cell(0, i)
        set_cell_bg(c, '0f766e')
        run_h = c.paragraphs[0].add_run(h)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(255, 255, 255)
    for i, (date, title, desc) in enumerate(revisions, 1):
        rev_tbl.cell(i, 0).text = date
        rev_tbl.cell(i, 1).text = title
        rev_tbl.cell(i, 2).text = desc
        if i % 2 == 0:
            set_cell_bg(rev_tbl.cell(i, 0), 'f0fdfa')
            set_cell_bg(rev_tbl.cell(i, 1), 'f0fdfa')
            set_cell_bg(rev_tbl.cell(i, 2), 'f0fdfa')

    add_heading_custom(doc, '2-5. 화면 목록', 2, '0f766e')
    doc.add_picture(screen_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3 = doc.add_paragraph('[ 그림 3 ] Cal AI 9개 화면 구성도')
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3.runs[0].font.size = Pt(9)
    cap3.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_page_break()

    # ── ③ 결과 및 배운점 ──
    add_heading_custom(doc, '③ 결과 및 배운점', 1)

    add_heading_custom(doc, '3-1. Before / After 비교', 2, '7c3aed')
    ba_tbl = doc.add_table(rows=2, cols=2)
    ba_tbl.style = 'Table Grid'
    ba_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_bg(ba_tbl.cell(0, 0), '1e3a5f')
    set_cell_bg(ba_tbl.cell(0, 1), '166534')
    h_b = ba_tbl.cell(0, 0).paragraphs[0].add_run('⬛ Before (도입 전)')
    h_b.font.bold = True
    h_b.font.color.rgb = RGBColor(255, 255, 255)
    h_a = ba_tbl.cell(0, 1).paragraphs[0].add_run('✅ After (도입 후)')
    h_a.font.bold = True
    h_a.font.color.rgb = RGBColor(255, 255, 255)
    ba_tbl.cell(1, 0).text = (
        '• 식단 기록: 앱 검색 + 수동 입력 10~15분\n'
        '• 개발: 코딩 지식 없어 아이디어를 앱으로 구현 불가\n'
        '• 반복 개선: 개발자에게 요청 → 수일 대기\n'
        '• QA: 별도 테스터 필요\n'
        '• 배포: 서버 지식 없어 퍼블리싱 불가'
    )
    ba_tbl.cell(1, 1).text = (
        '• 식단 기록: 사진 촬영 → 저장 10초 이내\n'
        '• calorie-ai-gamma.vercel.app 배포 완료\n'
        '  (누구나 스마트폰에서 접속·설치 가능)\n'
        '• Planner-Generator-Evaluator 자동 반복으로\n'
        '  하루 만에 9기능 SPA 완성\n'
        '• Evaluator가 자동 품질 검사 → 인간 QA 불필요\n'
        '• PWA 설치: Android "홈 화면에 추가" 지원'
    )

    doc.add_paragraph()
    add_heading_custom(doc, '3-2. 배운점 및 꿀팁', 2, '7c3aed')
    lessons = [
        '역할 분리가 핵심: "만드는 AI(Generator)와 평가하는 AI(Evaluator)를 분리"하면 자기 작업을 스스로 검증하는 편향을 방지할 수 있습니다. 같은 에이전트에게 만들고 평가하게 하면 셀프 칭찬이 발생합니다.',
        'PRD가 퀄리티를 결정한다: 첫 버전(PRD v1)보다 구체적으로 재작성한 PRD v2에서 생성된 결과물이 훨씬 완성도가 높았습니다. "AI에게 잘 시키는 것"은 결국 좋은 PRD/스펙 작성 능력입니다.',
        'evaluation_criteria.md를 Generator와 Evaluator가 공유: Generator는 채점 기준을 미리 읽고 구현하고, Evaluator는 동일 기준으로 채점합니다. 기준을 공유하지 않으면 서로 다른 기준으로 판단해 루프가 끝나지 않습니다.',
        'CHANGELOG.md로 히스토리 관리: 매 반복마다 변경 내용을 CHANGELOG에 자동 기록하면, 다음 세션에서 컨텍스트를 빠르게 복원할 수 있습니다.',
        '꿀팁 — Edge Runtime 전환: Vercel에서 @vercel/node 빌더 대신 export const config = { runtime: "edge" }를 사용하면 Cold Start 없이 빠른 응답이 가능합니다.',
    ]
    for tip in lessons:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tip)

    add_heading_custom(doc, '3-3. 아직 막힌 것 (도움 필요)', 2, '7c3aed')
    blockers = [
        '재분석 이미지 만료: meal_logs에 base64 data URL로 저장한 이미지는 DB 용량을 많이 차지합니다. Phase 2에서 Supabase Storage로 이전해야 하는데, 이미지 URL이 만료될 경우의 처리 방법이 아직 미결입니다.',
        'Supabase 무료 플랜 한계: 이미지를 base64로 저장하면 PostgreSQL 행 크기 제한에 걸릴 수 있습니다. 큰 이미지(2MB 이상)를 저장할 때의 처리 전략이 필요합니다.',
        '통계 화면 Phase 2: 주간/월간 SVG 바 차트는 구조만 만들어졌고, 실제 데이터 집계 쿼리(Supabase rpc 또는 그룹바이)와 연동이 필요합니다.',
        'Claude의 완벽하지 않은 Vercel 배포 이해: Edge Runtime 전환 과정에서 여러 번 시행착오가 있었습니다. 정확한 vercel.json 문법을 Claude가 항상 바로 맞추지는 못했습니다.',
    ]
    for b in blockers:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b)

    add_heading_custom(doc, '3-4. 실패 경험 (제일 값진 내용)', 2, '7c3aed')
    add_highlight_para(doc,
        '실패 1 — image_url null 하드코딩으로 재분석 기능 전체 파괴\n\n'
        'Generator R3에서 saveMeal() 함수 내부에 image_url: null을 하드코딩했습니다. '
        '당시 Evaluator도 이를 발견하지 못하고 합격 판정을 내렸습니다. '
        '결과적으로 "저장된 이미지로 재분석" 기능(FR-7)이 완전히 작동하지 않았습니다. '
        '원인은 Generator가 "지금 당장 이미지가 없는 케이스"만 테스트하고, '
        'Evaluator도 코드 정적 분석만 했기 때문입니다.\n\n'
        '교훈: 정적 코드 분석만으로는 런타임 버그를 잡을 수 없습니다. '
        'Evaluator에게 "실제 데이터 흐름 추적"을 지시하거나, '
        'Playwright MCP로 실제 브라우저 테스트를 추가해야 합니다.',
        'fff3cd'
    )
    doc.add_paragraph()
    add_highlight_para(doc,
        '실패 2 — Vercel Edge Runtime 전환 3번의 재시도\n\n'
        'vercel.json의 builds 배열과 @vercel/node 빌더를 사용하다가 '
        '"루트 server.js가 엔트리포인트로 오인"되는 문제가 발생했습니다. '
        'Claude가 제안한 첫 번째 수정(builds 배열 명시)도 일부만 작동했고, '
        '최종적으로는 builds 배열 제거 + outputDirectory + Edge Runtime 조합으로 해결했습니다. '
        '이 과정에서 3번의 배포-실패-수정 사이클이 필요했습니다.\n\n'
        '교훈: 배포 환경 설정은 "작동하는 최소 구성"부터 시작해서 점진적으로 추가해야 합니다. '
        '한 번에 많이 바꾸면 어느 설정이 문제인지 파악하기 어렵습니다.',
        'fff3cd'
    )
    doc.add_paragraph()
    add_highlight_para(doc,
        '실패 3 — 단일 프롬프트로 너무 많은 것을 요청\n\n'
        'R3에서 "온보딩 4단계 + 재분석 모달 + 통계 차트 + Fallback 입력 화면"을 한 번에 구현 요청했더니 '
        'Generator가 일부 기능(image_url 저장)을 놓쳤습니다. '
        '복잡한 요청일수록 Generator는 새 기능에 집중하면서 기존 기능을 퇴보시키는 경향이 있었습니다.\n\n'
        '교훈: 변경 범위를 작게 나눠서 요청하고, 각 반복마다 Evaluator가 회귀 테스트를 수행하도록 '
        'evaluation_criteria.md에 "기존 기능 유지 체크" 항목을 추가해야 합니다.',
        'fff3cd'
    )

    add_heading_custom(doc, '3-5. 앞으로의 적용 계획', 2, '7c3aed')
    plans = [
        'Phase 2 구현: 주간/월간 통계 차트(Supabase 그룹바이 쿼리 + SVG 바 차트), Supabase Storage 이미지 영구 저장, 연속 기록 스트릭 기능 추가',
        'Playwright MCP 테스트 자동화: Evaluator의 정적 코드 분석 한계를 보완하기 위해 실제 브라우저 시나리오 테스트(로그인 → 분석 → 저장 → 기록 확인)를 자동화',
        '업무 확장: 이번에 검증한 Planner-Generator-Evaluator 하네스 패턴을 다른 업무 자동화(주간 보고서 생성, 회의록 자동화)에도 적용 예정',
        'AI 정확도 개선: Gemini 프롬프트에 한국 음식 특화 예시(사진 + 영양소 데이터)를 프롬프트에 추가하여 한국 음식 인식 정확도 향상 시도',
        '팀원 공유: MBA 동기들에게 calorie-ai-gamma.vercel.app을 공유하고 실제 사용 피드백을 수집하여 다음 이터레이션에 반영',
    ]
    for plan in plans:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(plan)

    # ── 부록 ──
    doc.add_page_break()
    add_heading_custom(doc, '부록 — 배포 정보 및 체크리스트', 1)

    add_heading_custom(doc, '배포 URL', 2, '475569')
    url_p = doc.add_paragraph()
    run_url = url_p.add_run('https://calorie-ai-gamma.vercel.app')
    run_url.font.bold = True
    run_url.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)

    add_heading_custom(doc, 'Supabase 테이블 생성 SQL', 2, '475569')
    sql = doc.add_paragraph()
    sql.paragraph_format.left_indent = Cm(0.5)
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:fill'), 'f1f5f9')
    shd2.set(qn('w:val'), 'clear')
    pPr2 = sql._p.get_or_add_pPr()
    pPr2.append(shd2)
    run_sql = sql.add_run(
        "-- profiles 테이블\n"
        "CREATE TABLE profiles (\n"
        "  id UUID PRIMARY KEY REFERENCES auth.users ON DELETE CASCADE,\n"
        "  gender TEXT, age INTEGER, height_cm DECIMAL(5,1), weight_kg DECIMAL(5,1),\n"
        "  activity_level TEXT, goal_type TEXT,\n"
        "  daily_calorie_goal INTEGER DEFAULT 2000,\n"
        "  protein_goal INTEGER DEFAULT 150, carb_goal INTEGER DEFAULT 250,\n"
        "  fat_goal INTEGER DEFAULT 65, gemini_model TEXT DEFAULT 'flash'\n"
        ");\n"
        "ALTER TABLE profiles ENABLE ROW LEVEL SECURITY;\n\n"
        "-- meal_logs 테이블\n"
        "CREATE TABLE meal_logs (\n"
        "  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
        "  user_id UUID REFERENCES auth.users NOT NULL,\n"
        "  meal_type TEXT, food_name TEXT NOT NULL, calories INTEGER NOT NULL,\n"
        "  protein DECIMAL(5,1), carbs DECIMAL(5,1), fat DECIMAL(5,1),\n"
        "  image_url TEXT, notes TEXT, logged_at TIMESTAMPTZ DEFAULT NOW()\n"
        ");\n"
        "ALTER TABLE meal_logs ENABLE ROW LEVEL SECURITY;"
    )
    run_sql.font.size = Pt(8.5)
    run_sql.font.name = 'Courier New'

    add_heading_custom(doc, '제출 전 체크리스트', 2, '475569')
    checklist = [
        '① 소개 — 업무 선택 이유가 구체적으로 작성되었는가?',
        '② 진행 방법 — 스킬 설명과 프롬프트 원문이 포함되었는가?',
        '② 진행 방법 — 수정 과정(5회 이력)이 기록되었는가?',
        '③ 결과 — Before/After가 구체적 사례로 작성되었는가?',
        '③ 결과 — 실패 경험 3개가 포함되었는가?',
        '아키텍처 개요 이미지가 포함되었는가?',
        '기술 스택 이미지가 포함되었는가?',
        '화면 목록 이미지가 포함되었는가?',
        '민감정보(API 키)가 제거되었는가?',
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.add_run('☑ ' + item)

    output_path = os.path.join(OUT_DIR, 'MBA_중간점검_과제물_이현용.docx')
    doc.save(output_path)
    print(f'[OK] Word 문서 저장: {output_path}')
    return output_path


# ─────────────────────────────────────────────
# 실행
# ─────────────────────────────────────────────
if __name__ == '__main__':
    print('[*] 다이어그램 생성 중...')
    arch = make_architecture()
    tech = make_tech_stack()
    screen = make_screen_list()
    print('[*] Word 문서 생성 중...')
    result = create_word_doc(arch, tech, screen)
    print('\n[완료] 생성된 파일:')
    print(f'   - {arch}')
    print(f'   - {tech}')
    print(f'   - {screen}')
    print(f'   - {result}')
