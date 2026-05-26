# -*- coding: utf-8 -*-
"""Add GitHub to the tech stack diagram and clarify section 2-2."""
from __future__ import annotations

import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_DOC = BASE_DIR / "MBA_중간점검_과제물_이현용_부록2장씩.standard-docx"
OUTPUT_STANDARD = BASE_DIR / "MBA_중간점검_과제물_이현용_부록2장씩_GitHub추가.standard-docx"
OUTPUT_DOCX = BASE_DIR / "MBA_중간점검_과제물_이현용_부록2장씩_GitHub추가.docx"
TECH_PATH = BASE_DIR / "tech_stack_diagram.png"


TECH_PARAGRAPH = (
    "프론트엔드는 HTML5, CSS3, Vanilla JavaScript로 만든 모바일 우선 단일 페이지 앱(SPA)이다. "
    "사용자는 별도 설치 없이 브라우저에서 사진을 선택하고, 화면에서 분석 결과를 바로 확인한다. "
    "AI 분석은 서버 API가 Gemini 2.5 모델을 호출하는 방식으로 처리해 음식명, 예상 칼로리, 단백질, 탄수화물, 지방을 추정한다. "
    "로그인과 식단 기록 저장은 Supabase Auth와 PostgreSQL DB를 사용했다. "
    "작성한 코드는 GitHub 저장소에서 버전 관리하고, 변경 이력을 남겨 문제 발생 시 이전 상태를 확인할 수 있도록 했다. "
    "배포는 Vercel을 사용했으며 GitHub와 연동해 코드 변경 사항을 배포 흐름에 연결했다. "
    "Gemini API 키와 같은 민감 정보는 클라이언트 화면에 노출하지 않고 Vercel 서버 환경변수로 관리했다."
)


def setup_plot() -> None:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 180
    plt.rcParams["savefig.dpi"] = 280


def rounded_box(ax, x, y, w, h, text, face, edge=None, text_color="#0f172a", size=12, radius=0.08):
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0.04,rounding_size={radius}",
        linewidth=2,
        edgecolor=edge or face,
        facecolor=face,
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        color=text_color,
        fontsize=size,
        fontweight="bold",
        linespacing=1.25,
    )


def make_tech_stack() -> None:
    setup_plot()
    fig, ax = plt.subplots(figsize=(14.5, 8.4))
    fig.patch.set_facecolor("#ffffff")
    ax.set_xlim(0, 14.5)
    ax.set_ylim(0, 8.4)
    ax.axis("off")
    ax.text(7.25, 7.95, "Cal AI 기술 스택", ha="center", fontsize=24, fontweight="bold", color="#0f172a")

    layers = [
        (
            "사용자 화면",
            5.55,
            "#dbeafe",
            "#1d4ed8",
            [("HTML5", "구조"), ("CSS3", "모바일 UI"), ("Vanilla JS", "SPA 상태"), ("SVG/Canvas", "차트")],
        ),
        (
            "AI Agent 개발",
            3.25,
            "#dcfce7",
            "#15803d",
            [("Claude Code", "작업 조율"), ("Planner", "설계"), ("Generator", "구현"), ("Evaluator", "검증")],
        ),
        (
            "데이터와 배포",
            0.95,
            "#fef3c7",
            "#92400e",
            [
                ("Node API", "서버 호출"),
                ("Gemini", "음식 분석"),
                ("Supabase", "Auth + DB"),
                ("GitHub", "버전 관리"),
                ("Vercel", "배포 + 환경변수"),
            ],
        ),
    ]

    for title, y, bg, accent, items in layers:
        rounded_box(ax, 0.45, y, 13.6, 1.75, "", bg, edge="#cbd5e1")
        ax.text(0.85, y + 1.35, title, fontsize=14, fontweight="bold", color="#334155", ha="left")

        if len(items) == 5:
            item_w = 2.25
            gap = 0.28
            start_x = 1.05
        else:
            item_w = 2.35
            gap = 0.65
            start_x = 1.15

        for i, (name, desc) in enumerate(items):
            x = start_x + i * (item_w + gap)
            rounded_box(ax, x, y + 0.25, item_w, 0.82, f"{name}\n{desc}", accent, text_color="#ffffff", size=10.5)

    fig.savefig(TECH_PATH, bbox_inches="tight", facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)


def set_font(run, size=None, bold=None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run, size=10.5)


def replace_picture_after_heading(doc: Document, heading: str, image_path: Path, width: float) -> None:
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == heading)
    for paragraph in doc.paragraphs[start + 1 :]:
        if paragraph.text.strip().startswith(("2-3.", "2-4.", "3.", "부록.")):
            break
        if paragraph._element.xpath(".//w:drawing"):
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(width))
            return
    raise ValueError("Technology stack image paragraph not found.")


def update_doc() -> Path:
    make_tech_stack()

    doc = Document(SOURCE_DOC)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "2-2. 기술 스택":
            replace_paragraph_text(doc.paragraphs[i + 2], TECH_PARAGRAPH)
            break
    else:
        raise ValueError("Section heading not found: 2-2. 기술 스택")

    replace_picture_after_heading(doc, "2-2. 기술 스택", TECH_PATH, 6.3)

    doc.save(OUTPUT_STANDARD)
    shutil.copyfile(OUTPUT_STANDARD, OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(update_doc())
