# -*- coding: utf-8 -*-
"""Create the MBA midterm assignment report for the Cal AI project."""
from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent


def setup_plot() -> None:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def rounded_box(ax, xy, width, height, text, face, edge="#ffffff", color="#ffffff", size=10):
    x, y = xy
    box = FancyBboxPatch(
        (x, y),
        width,
        height,
        boxstyle="round,pad=0.08,rounding_size=0.08",
        linewidth=1.5,
        edgecolor=edge,
        facecolor=face,
    )
    ax.add_patch(box)
    ax.text(
        x + width / 2,
        y + height / 2,
        text,
        ha="center",
        va="center",
        fontsize=size,
        fontweight="bold",
        color=color,
        linespacing=1.3,
    )


def arrow(ax, start, end, color="#2563eb"):
    ax.annotate("", xy=end, xytext=start, arrowprops={"arrowstyle": "->", "lw": 2, "color": color})


def make_architecture() -> Path:
    setup_plot()
    fig, ax = plt.subplots(figsize=(12, 7))
    fig.patch.set_facecolor("#f8fafc")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.text(6, 6.55, "Cal AI 3-Agent 파이프라인 아키텍처", ha="center", fontsize=18, fontweight="bold", color="#0f172a")

    rounded_box(ax, (0.4, 3.35), 1.6, 0.75, "사용자\n요청", "#475569")
    rounded_box(ax, (2.5, 3.35), 1.8, 0.75, "Planner\nAgent", "#0284c7")
    rounded_box(ax, (5.0, 3.35), 1.8, 0.75, "Generator\nAgent", "#16a34a")
    rounded_box(ax, (7.5, 3.35), 1.8, 0.75, "Evaluator\nAgent", "#f59e0b")
    rounded_box(ax, (10.0, 3.35), 1.6, 0.75, "합격\n완료", "#22c55e")

    rounded_box(ax, (2.6, 4.65), 1.6, 0.55, "PRD.md", "#e0f2fe", edge="#0284c7", color="#075985", size=9)
    rounded_box(ax, (5.1, 4.65), 1.6, 0.55, "SPEC.md", "#dcfce7", edge="#16a34a", color="#166534", size=9)
    rounded_box(ax, (5.1, 2.15), 1.6, 0.55, "output/\nindex.html", "#f1f5f9", edge="#64748b", color="#334155", size=9)
    rounded_box(ax, (7.35, 2.15), 2.1, 0.55, "QA_REPORT.md\n점수/피드백", "#ffedd5", edge="#f59e0b", color="#9a3412", size=8)

    arrow(ax, (2.0, 3.72), (2.5, 3.72))
    arrow(ax, (4.3, 3.72), (5.0, 3.72))
    arrow(ax, (6.8, 3.72), (7.5, 3.72))
    arrow(ax, (9.3, 3.72), (10.0, 3.72))
    arrow(ax, (3.4, 4.65), (5.0, 4.02), "#64748b")
    arrow(ax, (5.9, 4.65), (5.9, 4.1), "#64748b")
    arrow(ax, (8.4, 3.35), (8.4, 2.7), "#64748b")
    arrow(ax, (7.4, 2.4), (5.9, 2.4), "#ef4444")
    arrow(ax, (5.9, 2.7), (5.9, 3.35), "#ef4444")
    ax.text(6.65, 2.05, "불합격 또는 조건부 합격 시 피드백 반영", ha="center", fontsize=9, color="#ef4444")

    y = 0.65
    for i, (label, face, desc) in enumerate(
        [
            ("HTML/CSS/JS", "#0ea5e9", "모바일 SPA"),
            ("Gemini 2.5", "#7c3aed", "이미지 분석"),
            ("Supabase", "#059669", "Auth + DB"),
            ("Vercel", "#111827", "배포 + API"),
        ]
    ):
        rounded_box(ax, (0.7 + i * 2.85, y), 2.2, 0.85, f"{label}\n{desc}", face, size=9)
    ax.plot([0.4, 11.6], [1.9, 1.9], "--", color="#cbd5e1")

    path = OUT_DIR / "architecture_diagram.png"
    fig.savefig(path, dpi=170, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def make_tech_stack() -> Path:
    setup_plot()
    fig, ax = plt.subplots(figsize=(11, 8))
    fig.patch.set_facecolor("#ffffff")
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.text(5.5, 7.55, "기술 스택 레이어 다이어그램", ha="center", fontsize=18, fontweight="bold", color="#0f172a")

    layers = [
        ("UI / 인터페이스", 5.55, "#dbeafe", [("HTML5", "구조"), ("CSS3", "모바일 UI"), ("Vanilla JS", "SPA 상태"), ("SVG", "차트")]),
        ("AI / Agent", 3.35, "#dcfce7", [("Claude Code", "오케스트레이션"), ("Planner", "설계"), ("Generator", "구현"), ("Evaluator", "검증")]),
        ("Data / Server", 1.15, "#fef3c7", [("Node API", "프록시"), ("Gemini", "분석"), ("Supabase", "Auth/DB"), ("Vercel", "배포")]),
    ]
    for title, y, bg, items in layers:
        rounded_box(ax, (0.35, y), 10.3, 1.65, "", bg, edge="#cbd5e1", color="#0f172a")
        ax.text(0.7, y + 1.35, title, fontsize=12, fontweight="bold", color="#334155")
        for i, (name, desc) in enumerate(items):
            x = 0.75 + i * 2.5
            rounded_box(ax, (x, y + 0.25), 2.05, 0.8, f"{name}\n{desc}", "#0f766e" if y == 3.35 else "#2563eb" if y == 5.55 else "#92400e", size=9)

    path = OUT_DIR / "tech_stack_diagram.png"
    fig.savefig(path, dpi=170, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def make_screen_list() -> Path:
    setup_plot()
    screens = [
        ("01", "온보딩", "성별/나이/키/체중\nTDEE 자동 계산", "#2563eb"),
        ("02", "인증", "이메일 회원가입\n게스트 모드", "#0284c7"),
        ("03", "홈", "오늘 칼로리 링\n식사 추가 FAB", "#16a34a"),
        ("04", "분석/수정", "Gemini 분석 결과\n사용자 보정 후 저장", "#22c55e"),
        ("05", "기록", "날짜별 식단\n수정/삭제/재분석", "#f59e0b"),
        ("06", "수동 입력", "AI 실패 시\n직접 식사 입력", "#ef4444"),
        ("07", "재분석", "저장 이미지 또는\n새 사진으로 재분석", "#8b5cf6"),
        ("08", "통계", "주/월 단위\n섭취 추세 차트", "#0f766e"),
        ("09", "설정", "목표/모델/프로필\n재설정", "#64748b"),
    ]
    fig, axes = plt.subplots(3, 3, figsize=(13, 9))
    fig.patch.set_facecolor("#f8fafc")
    fig.suptitle("Cal AI 화면 및 기능 목록", fontsize=18, fontweight="bold", color="#0f172a", y=0.98)
    for ax, (num, name, desc, color) in zip(axes.flatten(), screens):
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.axis("off")
        rounded_box(ax, (0.4, 0.45), 9.2, 9.0, "", "#ffffff", edge=color)
        circ = plt.Circle((1.55, 8.1), 0.8, color=color)
        ax.add_patch(circ)
        ax.text(1.55, 8.1, num, ha="center", va="center", color="#ffffff", fontweight="bold")
        ax.text(2.65, 8.1, name, ha="left", va="center", fontsize=13, fontweight="bold", color="#0f172a")
        ax.plot([1.0, 9.0], [6.7, 6.7], color=color, linewidth=1.5, alpha=0.35)
        ax.text(5.0, 4.05, desc, ha="center", va="center", fontsize=11, color="#334155", linespacing=1.5)

    path = OUT_DIR / "screen_list_diagram.png"
    fig.savefig(path, dpi=170, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color="0f172a"):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 12, True, color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, 10.5)


def add_prompt_box(doc, title, body):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_font(run, 10.5, True, "0f766e")
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Cm(0.4)
    box.paragraph_format.right_indent = Cm(0.2)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "f1f5f9")
    box._p.get_or_add_pPr().append(shd)
    run = box.add_run(body)
    set_font(run, 9)


def create_doc(arch, tech, screen) -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Agent 중간점검 과제물")
    set_font(run, 24, True, "0f172a")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Cal AI: 음식 사진 한 장으로 기록하는 AI 칼로리 트래커")
    set_font(run, 13, False, "16a34a")

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (k, v) in enumerate(
        [
            ("학생", "이현용"),
            ("날짜", dt.date.today().strftime("%Y년 %m월 %d일")),
            ("과목", "빅데이터와 인공지능"),
            ("과정", "인천대 MBA"),
        ]
    ):
        info.cell(row, 0).text = k
        info.cell(row, 1).text = v
        shade_cell(info.cell(row, 0), "e2e8f0")
    doc.add_page_break()

    add_heading(doc, "1. 소개", 1)
    add_heading(doc, "1-1. 적용 업무와 선택 이유", 2, "166534")
    doc.add_paragraph(
        "이번 과제에서 AI Agent를 적용한 업무는 '음식 사진 기반 칼로리 기록 앱' 개발이다. "
        "다이어트나 건강 관리를 시작한 사용자는 식단 기록이 중요하다는 것을 알지만, 실제로는 음식명을 검색하고 "
        "양을 추정해 매번 입력하는 과정이 번거로워 쉽게 중단한다. 그래서 사용자가 사진을 찍으면 Gemini가 음식명과 "
        "칼로리, 단백질, 탄수화물, 지방을 추정하고, 사용자는 결과를 한 화면에서 수정한 뒤 저장하는 흐름을 만들었다."
    )
    doc.add_paragraph(
        "이 주제를 선택한 이유는 MBA 수업에서 배운 AI Agent의 가치가 단순한 챗봇보다 '반복되는 업무 절차를 설계, "
        "실행, 검증하는 구조'에서 더 잘 드러난다고 보았기 때문이다. Planner가 요구사항을 SPEC으로 정리하고, "
        "Generator가 구현하며, Evaluator가 기준표에 따라 점검하는 구조를 적용하면 개발 경험이 부족한 사람도 "
        "실제 사용 가능한 MVP를 빠르게 만들 수 있다."
    )
    add_heading(doc, "1-2. 기존 문제와 기대 효과", 2, "166534")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Before: 기존 문제"
    table.cell(0, 1).text = "After: Agent 적용 기대 효과"
    shade_cell(table.cell(0, 0), "1e3a5f")
    shade_cell(table.cell(0, 1), "166534")
    table.cell(1, 0).text = "음식 검색과 수동 입력에 10분 이상 소요\n기록 피로도가 높아 지속성이 낮음\n기획, 구현, 검증이 한 프롬프트에 섞여 품질 관리가 어려움"
    table.cell(1, 1).text = "사진 선택 후 약 10초 내 분석 결과 확인\n수정 후 저장하는 간단한 기록 흐름\nPlanner-Generator-Evaluator 분리로 반복 개선 가능"
    add_heading(doc, "1-3. 전체 아키텍처 개요", 2, "166534")
    doc.add_picture(str(arch), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    add_heading(doc, "2. 진행 방법: 어떻게 만들었나", 1)
    add_heading(doc, "2-1. 만든 스킬과 에이전트 목록", 2, "0f766e")
    add_bullets(
        doc,
        [
            "CLAUDE.md: 전체 파이프라인을 조율하는 마스터 오케스트레이션 지침",
            "agents/planner.md: PRD를 읽고 화면, 데이터, 기능 요구사항을 SPEC.md로 구조화",
            "agents/generator.md: SPEC과 QA 피드백을 바탕으로 output/index.html 단일 SPA 구현",
            "agents/evaluator.md: evaluation_criteria.md 기준으로 기능성, UX, 기술 안정성, 완성도 점검",
            "agents/evaluation_criteria.md: Generator와 Evaluator가 공유하는 채점 기준 문서",
        ],
    )
    add_heading(doc, "2-2. 기술 스택", 2, "0f766e")
    doc.add_picture(str(tech), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "프론트엔드는 HTML/CSS/Vanilla JS 기반의 모바일 우선 SPA로 구성했다. AI 분석은 Gemini 2.5 Flash와 Pro를 "
        "서버 API를 통해 호출하고, 인증과 식단 기록은 Supabase Auth와 PostgreSQL을 사용했다. 배포는 Vercel을 사용했으며 "
        "API 키는 클라이언트에 노출하지 않고 서버 환경변수로 관리했다."
    )
    add_heading(doc, "2-3. 사용한 주요 프롬프트", 2, "0f766e")
    add_prompt_box(
        doc,
        "Planner 호출",
        "agents/planner.md와 evaluation_criteria.md를 읽고, PRD.md의 요구사항을 기반으로 칼로리 트래커 상세 화면 설계서를 SPEC.md로 작성하라.",
    )
    add_prompt_box(
        doc,
        "Generator 호출",
        "agents/generator.md를 읽고 SPEC.md 기반으로 Cal AI 전체를 output/index.html에 구현하라. Supabase JS CDN, Gemini API 호출, 모바일 Bottom Navigation을 포함하라.",
    )
    add_prompt_box(
        doc,
        "Evaluator 호출",
        "agents/evaluator.md와 evaluation_criteria.md를 읽고 SPEC.md와 output/index.html을 비교 검증하라. 기능성, UX, 기술 안정성, 완성도 점수를 QA_REPORT.md로 작성하라.",
    )
    add_prompt_box(
        doc,
        "배포 환경 수정",
        "Vercel에서 /api/config와 /api/analyze가 정상 인식되도록 vercel.json과 API 런타임 설정을 수정하고, Gemini 키는 서버 .env로 관리하라.",
    )

    add_heading(doc, "2-4. 수정 과정", 2, "0f766e")
    revisions = [
        ("R1", "초기 MVP 생성", "Auth, 홈, 기록, 설정 중심의 기본 SPA 생성. QA 8.7/10."),
        ("R2", "UX 피드백 반영", "회원가입 성공 메시지와 식사 카드 썸네일 개선. QA 9.2/10."),
        ("R3", "PRD v2 확장", "온보딩, TDEE, 재분석, 수동 입력 fallback, 통계 화면 추가. QA 8.1/10."),
        ("R4", "서버 프록시 전환", "Gemini API 키를 서버 .env로 이동, Supabase RLS 테이블 구성. QA 9.3/10."),
        ("R5", "UX 상용화 및 배포", "카메라/갤러리 바텀시트, 커스텀 모달, Vercel 배포 완료. QA 9.6/10."),
    ]
    rev = doc.add_table(rows=1 + len(revisions), cols=3)
    rev.style = "Table Grid"
    for i, h in enumerate(["차수", "작업", "결과"]):
        rev.cell(0, i).text = h
        shade_cell(rev.cell(0, i), "0f766e")
    for r, row in enumerate(revisions, 1):
        for c, val in enumerate(row):
            rev.cell(r, c).text = val
    add_heading(doc, "2-5. 화면 목록", 2, "0f766e")
    doc.add_picture(str(screen), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    add_heading(doc, "3. 결과 및 배운점", 1)
    add_heading(doc, "3-1. Before / After", 2, "7c3aed")
    ba = doc.add_table(rows=2, cols=2)
    ba.style = "Table Grid"
    ba.cell(0, 0).text = "Before"
    ba.cell(0, 1).text = "After"
    shade_cell(ba.cell(0, 0), "1e3a5f")
    shade_cell(ba.cell(0, 1), "166534")
    ba.cell(1, 0).text = "식단 기록은 검색과 수동 입력 중심이었다. 개발 과정도 아이디어, 구현, 검증이 섞여 있어 품질 개선 포인트를 분리하기 어려웠다."
    ba.cell(1, 1).text = "음식 사진 분석, 수동 보정, Supabase 저장, 재분석, 목표 설정, Vercel 배포까지 가능한 MVP가 완성되었다. 최종 QA 점수는 9.6/10이다."

    add_heading(doc, "3-2. 배운점과 꿀팁", 2, "7c3aed")
    add_bullets(
        doc,
        [
            "만드는 AI와 평가하는 AI를 분리하면 품질 관리가 훨씬 쉬워진다.",
            "좋은 PRD와 SPEC은 프롬프트보다 중요하다. 모호한 요구사항은 코드 품질의 변동으로 바로 이어진다.",
            "평가 기준을 Generator와 Evaluator가 함께 읽게 하면 목표가 일치한다.",
            "CHANGELOG와 QA_REPORT를 남기면 다음 반복에서 맥락 복원이 빠르다.",
            "API 키와 민감 정보는 초기에 서버 환경변수로 분리하는 편이 안전하다.",
        ],
    )
    add_heading(doc, "3-3. 아직 막힘 겪는 부분", 2, "7c3aed")
    add_bullets(
        doc,
        [
            "현재 이미지는 단기적으로 base64 형태를 활용하므로, 장기 운영에는 Supabase Storage 전환이 필요하다.",
            "통계 화면은 기본 집계 중심이라 주간/월간 분석과 개인화 인사이트는 Phase 2로 남아 있다.",
            "Vercel 런타임 설정은 여러 번 시행착오가 있었고, 배포 환경별 최소 설정을 별도 문서화할 필요가 있다.",
        ],
    )
    add_heading(doc, "3-4. 실패 경험", 2, "7c3aed")
    doc.add_paragraph(
        "가장 값진 실패는 R3에서 image_url을 null로 저장한 문제였다. 분석 단계에는 이미지 base64가 존재했지만 저장 단계에서 "
        "이를 meal_logs에 연결하지 않아 '저장된 이미지로 재분석' 기능이 비활성화되었다. 코드상으로는 앱이 정상 동작하는 것처럼 "
        "보였지만, 실제 사용자 시나리오에서는 핵심 기능 하나가 빠진 셈이었다. 이후 Evaluator가 단순 정적 점검뿐 아니라 데이터 "
        "흐름 관점에서 저장값과 후속 기능을 함께 검증해야 한다는 교훈을 얻었다."
    )
    doc.add_paragraph(
        "또 다른 실패는 Vercel 배포 설정이다. 처음에는 빌드 설정과 서버리스 함수 인식 방식이 맞지 않아 /api 경로가 정상 동작하지 "
        "않았다. 최종적으로는 API 파일 구조와 런타임 설정을 단순화해 해결했다. 배포 문제는 코드가 아니라 환경의 계약 문제이므로, "
        "작동하는 최소 구성부터 확인해야 한다는 점을 배웠다."
    )
    add_heading(doc, "3-5. 앞으로의 적용 계획", 2, "7c3aed")
    add_bullets(
        doc,
        [
            "Phase 2에서 Supabase Storage, 주간/월간 통계, 연속 기록 스트릭을 추가한다.",
            "Playwright 기반 실제 브라우저 테스트를 Evaluator 루프에 포함해 정적 분석의 한계를 보완한다.",
            "같은 Planner-Generator-Evaluator 구조를 주간 보고서 생성, 회의록 자동화, 데이터 분석 리포트 작성 업무에도 적용한다.",
            "MBA 동기들에게 배포 URL을 공유해 실제 사용 피드백을 수집하고 다음 반복에 반영한다.",
        ],
    )
    add_heading(doc, "부록. 배포 정보와 체크리스트", 1)
    doc.add_paragraph("배포 URL: https://calorie-ai-gamma.vercel.app")
    add_bullets(
        doc,
        [
            "소개, 진행 방법, 결과 및 배운점 섹션 작성 완료",
            "아키텍처, 기술 스택, 화면 목록 이미지 포함",
            "Before/After와 실패 경험 포함",
            "API 키 등 민감 정보는 문서에 포함하지 않음",
        ],
    )

    output = OUT_DIR / "MBA_중간점검_과제물_이현용.docx"
    temp_output = OUT_DIR / "MBA_중간점검_과제물_이현용.standard-docx"
    doc.save(temp_output)
    # The local PC may apply document DRM as soon as a file receives a .docx
    # extension. Keep the standard package copy above for automated validation.
    shutil.copyfile(temp_output, output)
    return output


if __name__ == "__main__":
    arch_path = make_architecture()
    tech_path = make_tech_stack()
    screen_path = make_screen_list()
    doc_path = create_doc(arch_path, tech_path, screen_path)
    print(doc_path)
