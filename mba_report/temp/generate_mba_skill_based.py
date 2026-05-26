# -*- coding: utf-8 -*-
"""
MBA 중간점검 과제물 — SKILL.md 기준 + 실제 프로젝트 데이터 기반 생성 스크립트

출처:
  - PRD.md          → 제품 개요, 기능 요구사항
  - QA_REPORT.md         → R1~R5 실제 QA 점수 및 이슈
  - CHANGELOG.md    → 수정 이력 (R1~R5 + 배포)
  - CLAUDE.md            → 실제 사용 프롬프트
  - agents/*.md          → 에이전트 역할 및 설계 기준
"""
from __future__ import annotations

import datetime as dt
import io
import re
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


BASE_DIR = Path(__file__).resolve().parent
ROOT_DIR = BASE_DIR.parent
PHOTO_DIR = BASE_DIR / "사진"
ARCH_PATH  = BASE_DIR / "temp" / "arch_skill.png"
TECH_PATH  = BASE_DIR / "temp" / "tech_skill.png"
SCREEN_PATH = BASE_DIR / "temp" / "screen_skill.png"
OUTPUT_DOCX = BASE_DIR / "MBA_중간점검_과제물_이현용_SKILL기반.docx"
OUTPUT_STD  = BASE_DIR / "MBA_중간점검_과제물_이현용_SKILL기반.standard-docx"

PHOTO_ORDER = [
    ("로그인",                          "로그인.png"),
    ("프로필 설정 - 나이, 성별, 무게, 키", "프로필 설정(나이_성별_무게_키).png"),
    ("프로필 설정 - 활동량",              "프로필 설정(활동량).png"),
    ("프로필 설정 - 체중 목표",           "프로필 설정(체중 목표).png"),
    ("프로필 설정 - 추천 칼로리 및 목표", "프로필 설정(추천 칼로리_목표 설정).png"),
    ("홈",                              "홈.png"),
    ("기록",                            "기록.png"),
    ("통계",                            "통계.png"),
    ("설정",                            "설정.png"),
]


# ─────────────────────────────────────────────────────────────────────────────
# Drawing helpers
# ─────────────────────────────────────────────────────────────────────────────

def setup_plot() -> None:
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def rbox(ax, x, y, w, h, text, face, edge=None, tc="#ffffff", fs=12, lw=2.0):
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.04,rounding_size=0.09",
        linewidth=lw, edgecolor=edge or face, facecolor=face, zorder=3,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text,
            ha="center", va="center", color=tc,
            fontsize=fs, fontweight="bold", linespacing=1.35, zorder=4)


def arr(ax, x0, y0, x1, y1, color="#2563eb", lw=2.2):
    ax.annotate("", xy=(x1, y1), xytext=(x0, y0),
                arrowprops={"arrowstyle": "->", "lw": lw, "color": color,
                            "shrinkA": 4, "shrinkB": 4}, zorder=5)


# ─────────────────────────────────────────────────────────────────────────────
# 1. Architecture diagram (개선판 — 겹침 없음)
# ─────────────────────────────────────────────────────────────────────────────

def make_architecture() -> Path:
    setup_plot()
    fig, ax = plt.subplots(figsize=(16, 11))
    fig.patch.set_facecolor("#f0f4f8")
    ax.set_xlim(0, 16);  ax.set_ylim(0, 11);  ax.axis("off")

    ax.text(8, 10.3, "Cal AI — 3-Agent 파이프라인 아키텍처",
            ha="center", fontsize=21, fontweight="bold", color="#0f172a", zorder=4)

    # Main pipeline ────────────────────────────────
    Y, H = 5.8, 1.3
    Cy = Y + H / 2  # 6.45
    mboxes = [
        (0.20, 1.80, "사용자\n요청",   "#475569"),
        (2.60, 2.30, "Planner\nAgent",  "#0284c7"),
        (5.80, 2.30, "Generator\nAgent","#16a34a"),
        (9.00, 2.30, "Evaluator\nAgent","#f59e0b"),
        (12.20,2.30, "합격\n완료",     "#22c55e"),
    ]
    for x, w, lbl, fc in mboxes:
        rbox(ax, x, Y, w, H, lbl, fc, fs=13)

    for x0, x1 in [(2.00,2.60),(4.90,5.80),(8.10,9.00),(11.30,12.20)]:
        arr(ax, x0, Cy, x1, Cy, "#2563eb", lw=2.4)

    # Doc boxes above ──────────────────────────────
    rbox(ax, 2.70, 7.8, 2.10, 0.75, "PRD.md",
         "#e0f2fe", edge="#0284c7", tc="#075985", fs=11, lw=1.8)
    rbox(ax, 5.90, 7.8, 2.10, 0.75, "SPEC.md",
         "#dcfce7", edge="#16a34a", tc="#166534", fs=11, lw=1.8)

    arr(ax, 3.75, 7.80, 3.75, Y + H, "#64748b", lw=2.0)  # PRD→Planner
    arr(ax, 6.95, 7.80, 6.95, Y + H, "#64748b", lw=2.0)  # SPEC→Generator
    arr(ax, 3.75, 8.55, 6.90, 8.55, "#64748b", lw=1.6)   # PRD→SPEC (Planner creates)
    ax.text(5.32, 8.68, "SPEC.md 생성", ha="center", fontsize=8.5,
            color="#475569", style="italic")

    # Feedback boxes below ─────────────────────────
    rbox(ax, 5.70, 3.3, 2.40, 0.85, "output/\nindex.html",
         "#f8fafc", edge="#64748b", tc="#334155", fs=10.5, lw=1.8)
    rbox(ax, 8.80, 3.3, 2.70, 0.85, "QA_REPORT.md\n점수 / 피드백",
         "#ffedd5", edge="#f59e0b", tc="#9a3412", fs=10, lw=1.8)

    arr(ax, 10.15, Y,     10.15, 4.15, "#64748b", lw=2.0)  # Evaluator→QA
    arr(ax, 8.80,  3.725,  8.10, 3.725,"#ef4444",  lw=2.2)  # QA→html
    arr(ax, 6.90,  4.15,   6.90, Y,    "#ef4444",  lw=2.2)  # html→Generator

    # Feedback label in highlight box
    bg = FancyBboxPatch((4.85, 2.32), 5.3, 0.42,
                        boxstyle="round,pad=0.04", linewidth=1.2,
                        edgecolor="#fca5a5", facecolor="#fff1f1", zorder=3)
    ax.add_patch(bg)
    ax.text(7.5, 2.53, "불합격 또는 조건부 합격 시 피드백 반영",
            ha="center", va="center", fontsize=10.5,
            color="#dc2626", fontweight="bold", zorder=5)

    # Divider & tech stack ─────────────────────────
    ax.plot([0.5, 15.5], [1.95, 1.95],
            linestyle=(0,(6,4)), color="#94a3b8", linewidth=1.8, zorder=2)

    tech = [("HTML/CSS/JS\n모바일 SPA","#0ea5e9"),
            ("Gemini 2.5\n이미지 분석","#7c3aed"),
            ("Supabase\nAuth + DB",    "#059669"),
            ("Vercel\n배포 + API",     "#111827")]
    tw, tg = 2.8, 0.55
    tx0 = (16 - (len(tech)*tw + (len(tech)-1)*tg)) / 2
    for i, (lbl, fc) in enumerate(tech):
        rbox(ax, tx0 + i*(tw+tg), 0.45, tw, 1.1, lbl, fc, fs=11)

    ARCH_PATH.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(ARCH_PATH, dpi=220, bbox_inches="tight",
                facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)
    return ARCH_PATH


# ─────────────────────────────────────────────────────────────────────────────
# 2. Tech stack diagram
# ─────────────────────────────────────────────────────────────────────────────

def make_tech_stack() -> Path:
    setup_plot()
    fig, ax = plt.subplots(figsize=(14, 8.5))
    fig.patch.set_facecolor("#ffffff")
    ax.set_xlim(0, 14);  ax.set_ylim(0, 8.5);  ax.axis("off")
    ax.text(7, 8.1, "Cal AI 기술 스택", ha="center",
            fontsize=22, fontweight="bold", color="#0f172a")

    layers = [
        ("사용자 화면", 5.5, "#dbeafe", "#1d4ed8",
         [("HTML5","구조"), ("CSS3","모바일 UI"), ("Vanilla JS","SPA 상태"), ("SVG","링·바 차트")]),
        ("AI Agent 개발", 3.2, "#dcfce7", "#15803d",
         [("Claude Code","오케스트레이션"), ("Planner","설계"), ("Generator","구현"), ("Evaluator","검증")]),
        ("데이터와 배포", 0.9, "#fef3c7", "#92400e",
         [("Node.js API","프록시"), ("Gemini 2.5","이미지 분석"), ("Supabase","Auth/DB"), ("Vercel","배포")]),
    ]
    for title, y, bg, accent, items in layers:
        rbox(ax, 0.4, y, 13.2, 1.85, "", bg, edge="#cbd5e1", tc="#0f172a", lw=1.5)
        ax.text(0.85, y+1.45, title, fontsize=13, fontweight="bold", color="#334155")
        for i, (nm, desc) in enumerate(items):
            rbox(ax, 1.0+i*3.15, y+0.28, 2.55, 0.95, f"{nm}\n{desc}",
                 accent, tc="#ffffff", fs=11)

    TECH_PATH.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(TECH_PATH, dpi=220, bbox_inches="tight",
                facecolor=fig.get_facecolor(), pad_inches=0.2)
    plt.close(fig)
    return TECH_PATH


# ─────────────────────────────────────────────────────────────────────────────
# 3. Screen list diagram
# ─────────────────────────────────────────────────────────────────────────────

def make_screen_list() -> Path:
    setup_plot()
    screens = [
        ("01","온보딩",   "나이/성별/키/체중\nHarris-Benedict TDEE","#2563eb"),
        ("02","인증",     "이메일 회원가입\n게스트 모드",           "#0284c7"),
        ("03","홈",       "오늘 칼로리 링차트\n식사 추가 FAB",      "#16a34a"),
        ("04","분석/수정","Gemini 분석 결과\n사용자 보정 후 저장",  "#22c55e"),
        ("05","기록",     "날짜별 식단\n수정/삭제/재분석",          "#f59e0b"),
        ("06","수동입력", "AI 실패 시\n직접 식사 입력(fallback)",   "#ef4444"),
        ("07","재분석",   "저장 이미지 또는\n새 사진으로 재분석",   "#8b5cf6"),
        ("08","통계",     "주간 평균\n목표 추세 SVG 바차트",        "#0f766e"),
        ("09","설정",     "목표/모델/프로필\n재설정",               "#64748b"),
    ]
    fig, axes = plt.subplots(3, 3, figsize=(15, 10.5))
    fig.patch.set_facecolor("#f8fafc")
    fig.suptitle("Cal AI 화면 및 기능 목록 (9개)",
                 fontsize=22, fontweight="bold", color="#0f172a", y=0.98)
    plt.subplots_adjust(hspace=0.25, wspace=0.2)

    for ax, (num, name, desc, color) in zip(axes.flatten(), screens):
        ax.set_xlim(0,10); ax.set_ylim(0,10); ax.axis("off")
        card = FancyBboxPatch((0.4,0.5),9.2,9.0,
                              boxstyle="round,pad=0.05", linewidth=2,
                              edgecolor=color, facecolor="#ffffff", zorder=2)
        band = FancyBboxPatch((0.4,8.0),9.2,1.5,
                              boxstyle="round,pad=0.05", linewidth=0,
                              edgecolor=color, facecolor=color, zorder=3)
        ax.add_patch(card); ax.add_patch(band)
        circ = plt.Circle((1.55,8.75), 0.6, color="#ffffff", zorder=4)
        ax.add_patch(circ)
        ax.text(1.55,8.75,num, ha="center",va="center",color=color,
                fontweight="bold",fontsize=11,zorder=5)
        ax.text(5.0,8.75,name, ha="center",va="center",
                fontsize=13.5,fontweight="bold",color="#ffffff",zorder=5)
        ax.text(5.0,4.2,desc, ha="center",va="center",
                fontsize=12,color="#334155",linespacing=1.6,zorder=4)

    SCREEN_PATH.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(SCREEN_PATH, dpi=220, bbox_inches="tight",
                facecolor=fig.get_facecolor(), pad_inches=0.15)
    plt.close(fig)
    return SCREEN_PATH


# ─────────────────────────────────────────────────────────────────────────────
# Word helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:   run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color:  run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color="0f172a"):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13, True, color)
    return p


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size)
    return p


def add_bullets(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size)


def add_prompt_box(doc, title, body):
    p = doc.add_paragraph()
    r = p.add_run(f"[{title}]")
    set_font(r, 10.5, True, "0f766e")
    box = doc.add_paragraph()
    box.paragraph_format.left_indent  = Cm(0.4)
    box.paragraph_format.right_indent = Cm(0.2)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "f1f5f9")
    box._p.get_or_add_pPr().append(shd)
    run = box.add_run(body)
    set_font(run, 9)


def set_cell_font(cell, size):
    for p in cell.paragraphs:
        for run in p.runs:
            set_font(run, size)


# ─────────────────────────────────────────────────────────────────────────────
# Build document — SKILL.md 형식 + 실제 프로젝트 데이터
# ─────────────────────────────────────────────────────────────────────────────
# PRD.md / CHANGELOG.md 파싱 헬퍼
# ─────────────────────────────────────────────────────────────────────────────

def _read_utf8(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        return ""


def _parse_prd(text: str) -> dict:
    out: dict = {}
    m = re.search(r'> (.+)', text)
    if m:
        out["description"] = m.group(1).strip()
    m = re.search(r'### 1\.2 해결하려는 문제\n\n((?:- .+\n?)+)', text)
    if m:
        out["problems"] = re.findall(r'- (.+)', m.group(1))
    m = re.search(r'### 1\.3 목표\n\n((?:\d+\. .+\n?)+)', text)
    if m:
        out["goals"] = re.findall(r'\d+\. (.+)', m.group(1))
    out["features"] = re.findall(
        r'\| (FR-\d+) \| ([^|]+?) \| ([^|]+?) \| (P\d) \|', text
    )
    return out


def _parse_changelog(text: str) -> list:
    entries = []
    current: dict | None = None
    in_impl = False
    for line in text.split("\n"):
        m = re.match(r'^## \[(\d{4}-\d{2}-\d{2})\] (.+)', line)
        if m:
            if current:
                entries.append(current)
            current = {"date": m.group(1), "title": m.group(2), "impl": [], "qa": "-"}
            in_impl = False
        elif current:
            if line.startswith("### 구현 내용"):
                in_impl = True
            elif line.startswith("###"):
                in_impl = False
            elif in_impl and line.startswith("- "):
                current["impl"].append(line[2:].strip())
            qa_m = re.search(r'최종 점수: ([\d.]+/10)', line)
            if qa_m:
                current["qa"] = qa_m.group(1)
    if current:
        entries.append(current)
    return entries


# ─────────────────────────────────────────────────────────────────────────────

def create_doc(arch: Path, tech: Path, screen: Path) -> Path:
    prd_text = _read_utf8(ROOT_DIR / "PRD.md")
    prd = _parse_prd(prd_text)
    cl_entries = _parse_changelog(_read_utf8(ROOT_DIR / "CHANGELOG.md"))

    doc = Document()

    # Global normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    # ── Cover ──────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("AI Agent 중간점검 과제물")
    set_font(r, 24, True, "0f172a")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run("Cal AI: 음식 사진 한 장으로 기록하는 AI 칼로리 트래커")
    set_font(r, 13, False, "16a34a")

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (k, v) in enumerate([
        ("학생",  "이현용"),
        ("날짜",  dt.date.today().strftime("%Y년 %m월 %d일")),
        ("과목",  "빅데이터와 인공지능"),
        ("과정",  "인천대 MBA"),
    ]):
        info.cell(row, 0).text = k
        info.cell(row, 1).text = v
        shade_cell(info.cell(row, 0), "e2e8f0")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # ① 소개
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "① 소개", 1)

    # 1-1. 적용 업무와 선택 이유
    add_heading(doc, "1-1. 적용 업무와 선택 이유", 2, "166534")
    add_para(doc,
        "이번 과제에서 AI Agent를 적용한 업무는 '음식 사진 기반 칼로리 기록 앱' 개발이다. "
        "사용자가 스마트폰으로 음식 사진을 찍으면 Google Gemini 2.5 AI가 음식명·칼로리·단백질·탄수화물·지방을 "
        "자동 추정하고, 사용자는 결과를 한 화면에서 보정한 뒤 저장하는 흐름으로 구성했다. "
        "Harris-Benedict 공식 기반 TDEE 자동 계산, 날짜별 식단 히스토리, SVG 링 차트 대시보드, "
        "Supabase 영구 저장, Vercel 배포까지 갖춘 완성형 모바일 웹 앱(PWA)이다. "
        "배포 URL: https://calorie-ai-gamma.vercel.app"
    )
    # 개인 다이어트 추가 문단
    add_para(doc,
        "다만, 이번 과제에서 선택한 주제는 현재 내가 수행하는 실제 업무와 직접적인 연관이 있는 주제는 아니다. "
        "개인적으로 진행 중인 다이어트와 건강 관리 습관을 개선하는 데 도움을 받기 위해 이 주제를 선택했다. "
        "매일 식단을 기록하고 칼로리를 확인하는 과정에서 꼭 필요한 기능만 자동화하면, "
        "MBA 수업에서 배운 AI Agent 접근 방법을 개인 생활의 문제 해결에도 충분히 적용할 수 있다고 판단했다."
    )
    # (삭제) "이 주제를 선택한 이유는 MBA 수업에서 배운 AI Agent의 가치가..." 문단은 포함하지 않음

    # 1-2. 기존 문제와 기대 효과
    add_heading(doc, "1-2. 기존 문제와 기대 효과", 2, "166534")
    # 기존 Before/After 테이블 삭제 → 아래 새 내용으로 교체
    add_para(doc,
        "이 주제를 선택한 이유는 기존 칼로리 계산기 앱에는 내가 실제로 사용하지 않는 기능이 지나치게 많이 "
        "포함되어 있고, 월 사용 비용도 약 4,000원에서 12,000원 수준으로 부담이 크기 때문이다. "
        "그래서 실제로 내가 필요한 칼로리 계산 기능과 식단 기록 기능만 구현하고, "
        "기타 불필요한 기능은 제외한 자동화 앱으로 직접 설계했다. "
        "실제 사용 비용도 이미지 분석 기준으로 한 달에 약 3장, 약 400원 정도가 소요될 것으로 예상한다."
    )

    # PRD 기반 해결 방향 요약
    add_para(doc,
        "PRD v2.1에서 정의한 해결 방향은 세 가지다. "
        "첫째, 음식 기록 시간을 10초 이내로 줄인다(사진 촬영 → 저장까지 3탭 이내). "
        "둘째, AI 분석 후 수정·저장을 한 화면에서 완결한다. "
        "셋째, Harris-Benedict 공식으로 개인 맞춤 TDEE를 자동 계산해 목표 대비 섭취량을 직관적으로 보여준다."
    )

    # 1-3. 전체 아키텍처 개요
    add_heading(doc, "1-3. 전체 아키텍처 개요", 2, "166534")
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(arch), width=Inches(6.3))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # ② 진행 방법 — 어떻게 만들었나
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "② 진행 방법 — 어떻게 만들었나", 1)

    # 2-1. 만든 스킬 목록
    add_heading(doc, "2-1. 만든 스킬(에이전트) 목록", 2, "0f766e")
    add_para(doc,
        "이번 프로젝트는 Claude Code CLI를 오케스트레이터로 활용하는 3-Agent 하네스 구조로 설계했다. "
        "핵심은 '만드는 AI'와 '평가하는 AI'를 완전히 분리하는 것이다."
    )
    add_bullets(doc, [
        "CLAUDE.md — 마스터 오케스트레이터: Planner → Generator → Evaluator 파이프라인 흐름 전체를 지시. "
          "불합격·조건부 시 Generator를 최대 3회 재호출하는 루프 로직 포함.",
        "agents/planner.md — Planner Agent: PRD.md를 읽고 SPEC.md(화면별 UI 명세 + 데이터 연동 상세)를 생성. "
          "구현 가능성을 고려해 HTML/CSS/JS 단일 파일 범위 내 설계.",
        "agents/generator.md — Generator Agent: SPEC.md와 QA_REPORT.md의 개선 지시를 읽고 "
          "output/index.html(SPA 단일 파일)을 구현. 자체 점검 문서(SELF_CHECK.md)도 작성.",
        "agents/evaluator.md — Evaluator Agent: evaluation_criteria.md 채점 기준으로 "
          "기능성(40%)·UX/디자인(30%)·기술 품질(20%)·기능 완성도(10%)를 점검. QA_REPORT.md 생성.",
        "agents/evaluation_criteria.md — 공유 채점 기준: Generator와 Evaluator가 동일한 기준을 읽어 "
          "목표를 일치시키는 핵심 문서. 합격(7.0+)·조건부·불합격 판정 기준 포함.",
    ])

    # 2-2. 기술 스택
    add_heading(doc, "2-2. 기술 스택", 2, "0f766e")
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.add_run().add_picture(str(tech), width=Inches(6.3))
    add_para(doc,
        "프론트엔드는 HTML/CSS/Vanilla JS 기반 모바일 우선 SPA(max-width 430px)로 구성했다. "
        "AI 이미지 분석은 Google Gemini 2.5 Flash(기본, 빠름·저비용)와 Pro(고정밀, 설정에서 전환 가능)를 "
        "Node.js 서버 API(/api/analyze)를 통해 호출한다. "
        "인증과 식단 기록 영구 저장에는 Supabase Auth + PostgreSQL(RLS 적용)을 사용하고, "
        "배포는 Vercel Edge Runtime으로 처리한다. API 키는 모두 서버 .env로 관리해 클라이언트에 노출되지 않는다."
    )

    # 2-3. 사용한 실제 프롬프트 (PRD.md 기반)
    add_heading(doc, "2-3. 사용한 실제 프롬프트 (PRD.md 기반)", 2, "0f766e")
    add_para(doc,
        "아래는 PRD.md(제품 요구사항 정의서 v2.1)의 핵심 내용이다. "
        "이 문서가 Planner 에이전트의 입력값이 되어 SPEC.md 생성 → Generator → Evaluator "
        "파이프라인을 구동했다."
    )

    # 제품 한줄 설명
    if prd.get("description"):
        add_prompt_box(doc, "1.1 제품 한줄 설명", prd["description"])

    # 해결하려는 문제
    if prd.get("problems"):
        add_prompt_box(doc, "1.2 해결하려는 문제",
                       "\n".join(f"• {p}" for p in prd["problems"]))

    # 목표
    if prd.get("goals"):
        add_prompt_box(doc, "1.3 목표",
                       "\n".join(f"{i+1}. {g}" for i, g in enumerate(prd["goals"])))

    # 기능 요구사항 테이블
    if prd.get("features"):
        add_heading(doc, "2.x 기능 요구사항 (FR)", 3)
        fr_tbl = doc.add_table(rows=1 + len(prd["features"]), cols=4)
        fr_tbl.style = "Table Grid"
        for ci, h in enumerate(["FR", "기능명", "설명", "우선순위"]):
            fr_tbl.cell(0, ci).text = h
            shade_cell(fr_tbl.cell(0, ci), "0f766e")
            for run in fr_tbl.cell(0, ci).paragraphs[0].runs:
                set_font(run, 9, True, "ffffff")
        for ri, (fr, name, desc, prio) in enumerate(prd["features"], 1):
            for ci, val in enumerate([fr, name.strip(), desc.strip(), prio]):
                fr_tbl.cell(ri, ci).text = val
                set_cell_font(fr_tbl.cell(ri, ci), 9)

    # 2-4. 수정 과정 (CHANGELOG.md 기반)
    add_heading(doc, "2-4. 수정 과정 (CHANGELOG.md 기반)", 2, "0f766e")
    add_para(doc,
        "아래는 CHANGELOG.md에 기록된 실제 수정 이력이다. "
        "CLAUDE.md 오케스트레이터가 Evaluator 판정을 읽고 합격 전까지 Generator를 재호출하며 누적된 기록이다."
    )
    if cl_entries:
        rev = doc.add_table(rows=1 + len(cl_entries), cols=4)
        rev.style = "Table Grid"
        for ci, h in enumerate(["날짜", "변경 제목", "주요 구현 내용", "QA 점수"]):
            rev.cell(0, ci).text = h
            shade_cell(rev.cell(0, ci), "0f766e")
            for run in rev.cell(0, ci).paragraphs[0].runs:
                set_font(run, 9.5, True, "ffffff")
        for ri, entry in enumerate(cl_entries, 1):
            impl_summary = "; ".join(entry["impl"][:2])
            if len(entry["impl"]) > 2:
                impl_summary += f" 외 {len(entry['impl'])-2}건"
            for ci, val in enumerate([
                entry["date"],
                entry["title"][:45],
                impl_summary[:120],
                entry["qa"],
            ]):
                rev.cell(ri, ci).text = val
                set_cell_font(rev.cell(ri, ci), 9)
    else:
        add_para(doc, "[CHANGELOG.md 파일을 읽을 수 없음]")

    # 2-5. 화면 목록
    add_heading(doc, "2-5. 화면 목록 (9개 화면, PRD v2 기준)", 2, "0f766e")
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run().add_picture(str(screen), width=Inches(6.3))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # ③ 결과 및 배운점
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "③ 결과 및 배운점", 1)

    # 3-1. 배운점과 꿀팁
    add_heading(doc, "3-1. 배운점과 꿀팁", 2, "7c3aed")
    add_bullets(doc, [
        "만드는 AI(Generator)와 평가하는 AI(Evaluator)를 분리하면 품질 관리가 체계화된다. "
          "같은 LLM이 만들고 평가하면 자기 오류를 잘 못 찾는다.",
        "PRD와 SPEC 문서의 품질이 프롬프트보다 중요하다. "
          "모호한 요구사항은 Generator 코드 품질의 변동으로 직결된다 (R3 점수 하락 경험).",
        "evaluation_criteria.md(채점 기준)를 Generator와 Evaluator가 동시에 읽게 하면 목표가 일치된다. "
          "목표가 다르면 아무리 반복해도 점수가 오르지 않는다.",
        "CHANGELOG와 QA_REPORT를 누적 기록하면 다음 반복에서 맥락 복원 시간이 크게 줄어든다. "
          "AI는 맥락이 길수록 판단이 정확해진다.",
        "API 키와 민감 정보는 프로젝트 초기에 서버 환경변수로 분리해야 한다. "
          "나중에 분리하려면 코드 전체를 수정해야 한다.",
        "Vercel 배포 문제는 코드 오류가 아니라 '환경의 계약 문제'다. "
          "작동하는 최소 설정부터 확인하는 것이 시간을 아낀다.",
    ])

    # 3-2. 아직 막힌 것
    add_heading(doc, "3-2. 아직 막힌 것 (도움 필요한 부분)", 2, "7c3aed")
    add_bullets(doc, [
        "현재 이미지는 base64로 meal_logs에 직접 저장한다. DB 용량 효율을 위해 "
          "Supabase Storage로 전환이 필요하지만, Phase 2로 미뤄진 상태다.",
        "통계 화면 SVG 바 차트는 구현됐지만, 주간/월간 집계 및 개인화 인사이트 기능은 아직 Phase 2다.",
        "Vercel Edge Runtime에서 긴 이미지 데이터(base64)를 처리할 때 Payload 크기 제한 이슈가 잠재적으로 존재한다.",
        "PWA 서비스 워커 캐싱 전략이 기본 수준이라, 오프라인 상태에서의 완전한 동작은 아직 보장되지 않는다.",
    ])

    # 3-3. 실패 경험 (실제 QA_REPORT 기반)
    add_heading(doc, "3-3. 실패 경험 (제일 값진 내용)", 2, "7c3aed")
    add_para(doc,
        "[실패 1] image_url null 하드코딩 (R3 — 기능성 7/10 원인)\n"
        "R3에서 Gemini 분석 후 saveMeal() 함수가 meal_logs에 항상 image_url: null을 "
        "저장하는 버그가 있었다. 분석 단계에서 currentBase64 변수에 이미지 데이터가 있었지만 "
        "INSERT 시 연결하지 않아, '저장된 이미지로 재분석' 버튼이 항상 비활성화되는 문제가 발생했다. "
        "코드상으로는 앱이 정상 작동하는 것처럼 보였지만, 실제 사용 시나리오에서는 핵심 기능 하나가 "
        "빠진 셈이었다. Evaluator가 단순 정적 점검뿐 아니라 데이터 흐름 관점에서 "
        "저장값과 후속 기능을 함께 검증해야 한다는 교훈을 얻었다."
    )
    add_para(doc,
        "[실패 2] Vercel 배포 설정 오류 (R4·R5 사이)\n"
        "vercel.json에 builds 배열을 명시했을 때 루트의 server.js가 엔트리포인트로 오인되어 "
        "/api 경로가 작동하지 않았다. 최종적으로 builds 배열을 제거하고 "
        "outputDirectory: 'output'과 Edge Runtime 설정으로 단순화해 해결했다. "
        "배포 문제는 코드 로직이 아니라 플랫폼 환경의 인식 방식 문제이므로, "
        "작동하는 최소 설정(Minimal Config)부터 확인하는 습관이 필요하다."
    )
    add_para(doc,
        "[실패 3] R3에서 QA 점수 하락 경험 (9.2 → 8.1)\n"
        "R2에서 9.2/10을 받은 뒤 PRD v2로 기능을 대폭 확장한 R3가 8.1/10으로 오히려 낮아졌다. "
        "새 기능이 많아질수록 QA 기준을 충족하지 못하는 부분이 늘어났기 때문이다. "
        "기능 추가와 품질 유지는 균형이 필요하며, 확장 후에는 반드시 회귀 검증이 필요하다는 점을 배웠다."
    )

    # 3-4. 앞으로의 적용 계획
    add_heading(doc, "3-4. 앞으로의 적용 계획", 2, "7c3aed")
    add_bullets(doc, [
        "Phase 2: Supabase Storage로 이미지 영구 저장 전환, 주간/월간 통계 차트, 연속 기록 스트릭.",
        "Playwright 기반 실제 브라우저 E2E 테스트를 Evaluator 루프에 포함해 정적 분석의 한계를 보완.",
        "동일한 Planner-Generator-Evaluator 파이프라인 구조를 주간 업무 보고서 자동 생성, "
          "회의록 요약, 데이터 분석 리포트 작성 등 실제 업무 자동화에도 적용할 예정.",
        "MBA 동기들에게 배포 URL(calorie-ai-gamma.vercel.app)을 공유해 실사용 피드백을 수집하고 "
          "다음 반복 개선에 반영.",
    ])

    # ══════════════════════════════════════════════════════════════════════
    # 부록 A. 배포 정보 및 체크리스트
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "부록 A. 배포 정보 및 체크리스트", 1)
    add_para(doc, "배포 URL: https://calorie-ai-gamma.vercel.app")
    add_para(doc, "GitHub:   https://github.com/yong9098-design/calorie_AI")
    add_bullets(doc, [
        "① 소개, ② 진행 방법, ③ 결과 및 배운점 섹션 완성",
        "아키텍처·기술 스택·화면 목록 이미지 포함",
        "실패 경험 포함",
        "PRD.md 기반 기능 요구사항 원문 포함",
        "CHANGELOG.md 기반 수정 이력 반영",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 부록 B. 실제 화면 캡처 (2장씩)
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "부록 B. 실제 화면 캡처", 1)
    add_para(doc,
        "아래 이미지는 Cal AI 앱의 실제 사용 흐름에 따라 정리했으며, 한 페이지에 두 장씩 배치하였다."
    )

    def remove_borders(tbl):
        tpr = tbl._tbl.tblPr
        brd = tpr.first_child_found_in("w:tblBorders")
        if brd is None:
            brd = OxmlElement("w:tblBorders"); tpr.append(brd)
        for edge in ("top","left","bottom","right","insideH","insideV"):
            el = brd.find(qn(f"w:{edge}"))
            if el is None:
                el = OxmlElement(f"w:{edge}"); brd.append(el)
            el.set(qn("w:val"), "nil")

    def photo_w(path: Path) -> Inches:
        try:
            with Image.open(path) as img:
                w, h = img.size
            return Inches(2.72) if h > w else Inches(5.5)
        except Exception:
            return Inches(2.72)  # 모바일 스크린샷은 세로형 기본값

    for pair_start in range(0, len(PHOTO_ORDER), 2):
        if pair_start > 0:
            doc.add_page_break()
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        remove_borders(tbl)
        for col in range(2):
            idx = pair_start + col
            if idx >= len(PHOTO_ORDER):
                break
            caption, filename = PHOTO_ORDER[idx]
            cell = tbl.cell(0, col)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.width = Inches(3.1)
            cell.text = ""
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(4)
            set_font(cp.add_run(f"부록 B-{idx+1}. {caption}"), 9.5, True)
            ip = cell.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            photo_path = PHOTO_DIR / filename
            try:
                img_bytes = photo_path.read_bytes()
                ip.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(2.72))
            except Exception:
                set_font(ip.add_run("[사진 없음]"), 8)
            fnp = cell.add_paragraph()
            fnp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(fnp.add_run(filename), 7.5)

    OUTPUT_STD.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_STD)
    shutil.copyfile(OUTPUT_STD, OUTPUT_DOCX)
    return OUTPUT_DOCX


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("[1] 아키텍처 다이어그램 생성 중...")
    arch = make_architecture()
    print(f"   -> {arch}")

    print("[2] 기술 스택 다이어그램 생성 중...")
    tech = make_tech_stack()
    print(f"   -> {tech}")

    print("[3] 화면 목록 다이어그램 생성 중...")
    screen = make_screen_list()
    print(f"   -> {screen}")

    print("[4] Word 문서 생성 중...")
    docx = create_doc(arch, tech, screen)
    print(f"\n완료: {docx}")
